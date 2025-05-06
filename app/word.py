from __future__ import division
#from trace import _FileModuleFunction
from docx import Document
from docx.shared import Inches
from pymongo import MongoClient
import sys
import os
import io
import uuid
import boto3
from app.config import Config
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,RGBColor,Cm
from docx.enum.text import WD_LINE_SPACING,WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from zappa.asynchronous import task
from flask import send_file, jsonify
from docx.enum.table import WD_ALIGN_VERTICAL
from pymongo.collation import Collation
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
import itertools
import math
import time
import certifi


s3_client = boto3.client('s3')


global globaltotalLines
global totalLinesAvailableOne
global totalLinesAvailableTwo
global titleTotalLength
global subtitleTotalLength
global entryTotalLength
global actualColumn
global lastTitle

globaltotalLines=60
totalLinesAvailableOne=60
totalLinesAvailableTwo=60
titleTotalLength=48
subtitleTotalLength=37
entryTotalLength=44
actualColumn="One"
lastTitle=""

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = '- '
    page_run._r.append(t1)

    font = page_run.font
    font.name="Arial"
    font.size= Pt(8)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    
    font = page_num_run.font
    font.name="Arial"
    font.size= Pt(8)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' -'
    of_run._r.append(t2)

    font = of_run.font
    font.name="Arial"
    font.size= Pt(8)
    

def add_hyperlink1(paragraph,text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url,RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text =text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    r.font.bold=True 

    return hyperlink

def add_hyperlink(paragraph,text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url,RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text =text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def generateWordDocITPSOR(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):

    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection}).sort("sortkey1",1)

    # Creation of the word document

    document = Document()

    # Implementing the "Two columns display" feature

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # Marging of the document

    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
   
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Adding the header to the document
    
    header=document.sections[0].header
    
    ################## SUBHEADER ###############################################
    
    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_sub_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    ################## SORENTRY ###############################################
    
    stlSorentry = document.styles.add_style('sorentry', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.bold = True
    
    pfSorentry = stlSorentry.paragraph_format

    # Line spacing
    
    pfSorentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
        
    ################## SORNORM ###############################################
    
    stlSornorm = document.styles.add_style('sornorm', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
   
    stlSornormFont=stlSornorm.font
    stlSornormFont.name = 'Arial'
    stlSornormFont.size = Pt(8)
    
    pfSornorm = stlSornorm.paragraph_format

    # Indentation
    
    pfSornorm.left_indent = Inches(0.50)

    # Line spacing
    
    pfSornorm.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    ################## SORNOTE ###############################################
    
    stlSornote = document.styles.add_style('sornote', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSornoteFont=stlSornote.font
    stlSornoteFont.name = 'Arial'
    stlSornoteFont.size = Pt(8)    
    
    pfSornote = stlSornote.paragraph_format
    
    # Indentation
    
    pfSornote.left_indent = Inches(0.68)

    # Line spacing
    
    pfSornote.line_spacing_rule =  WD_LINE_SPACING.SINGLE
     
    ################## WRITING THE DOCUMENT ###############################################
    
    # Adding the Header to the document
    
    p=header.add_paragraph(myTitle.upper(), style='New Heading')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adding the sub Header to the document
    
    p1=header.add_paragraph(paramSubTitle, style='New sub Heading')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = p1.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(30)

    p3=header.add_paragraph("", style='New sub Heading')

    # Adding the Footer
    
    #footer = document.sections[0].footer
    #paragraph=footer.add_paragraph("No footer defined",style="New Heading")

    myRecords=setOfData

    for record in myRecords:
        print(record)
        try :
            sorentry= record['sorentry'].strip()
        except :
            sorentry=""
        
        try :
            if record['docsymbol'].strip()!="":
                docSymbol=record['docsymbol'].strip()
            else:
                docSymbol="Symbol not used."
        except :
            docSymbol="Symbol not used."

        try :
            sornorm=record['sornorm'].strip()
        except :
            sornorm=""

        try :
            sornote=record['sornote'].strip()
        except :
            sornote=""

        # Adding sorentry content
        p=document.add_paragraph(sorentry,style=stlSorentry)
        
        # Breaks management
        paragraph_format = p.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

        
        # Breaks management
        paragraph_format.keep_together = True
        paragraph_format.keep_with_next = True
        
        # Force a tab
        p.add_run("\t")
        
        # Create the hyperlink for the document symbol
        if docSymbol!="Symbol not used." :
            if (bodysession[0] in ["E"]):
                recup=docSymbol.split(" ")
                add_hyperlink(p,docSymbol,Config.url_prefix+recup[0])
            else:
                add_hyperlink(p,docSymbol,Config.url_prefix+docSymbol)
        else:
            x=p.add_run("Symbol not used.")
            x.bold=False
            x.italic=True
            
            # Delete further space generated
            
            x_format = p.paragraph_format
            x_format.space_before = Pt(0)
            x_format.space_after = Pt(0)
        
        if sornorm!="" :
            # Adding sornorm content           
            p2=document.add_paragraph(sornorm,style=stlSornorm)
            
            # Breaks management
            paragraph_format = p2.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)

            # Breaks management
            paragraph_format.keep_together = True
            paragraph_format.keep_with_next = True

        if sornote!="" :
            # Adding sornote content  
            p3=document.add_paragraph(sornote,style=stlSornote)
            
            # Breaks management
            paragraph_format = p3.paragraph_format
            paragraph_format.space_before = Pt(0)

            # Breaks management
            paragraph_format.keep_together = True
            
        if sornote=="" and sornorm=="" :
            paragraph_format = p.paragraph_format
            paragraph_format.line_spacing =  1.5

    add_page_number(document.sections[0].footer.paragraphs[0])
    return document

def generateWordDocITPITSC(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle=paramSubTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})
    # Creation of the word document
    # Creation of the word document
    document = Document()

    # Two columns display

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # Marging of the document

    section.top_margin = Inches(1)
    section.header_distance=Inches(0.5)
    section.bottom_margin = Inches(1)
    section.footer_distance=Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1) 

    ################## HEADER ###############################################

    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']

    # Font settings

    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    # Adding the header to the document

    header=document.sections[0].header

    ################## SUBHEADER ###############################################

    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_sub_heading_style.base_style = styles['Heading 1']

    # Font settings

    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    ################## itshead ###############################################

    stlItsHead = document.styles.add_style('itshead', WD_STYLE_TYPE.PARAGRAPH)
    stlItsHead.quick_style =True

    # Font name

    stlItsHeadFont=stlItsHead.font
    stlItsHeadFont.name = 'Arial'
    stlItsHeadFont.size = Pt(9)
    stlItsHeadFont.bold = True

    pfItsHead = stlItsHead.paragraph_format
    pfItsHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsHead.space_before=Pt(8)
    pfItsHead.space_after=Pt(3)
    pfItsHead.keep_together = True

    ################## column break style #####################################
    stl_its_col_break = document.styles.add_style('itsColBreak', WD_STYLE_TYPE.PARAGRAPH)
    stl_its_col_break.quick_style =True


    stl_its_col_break_font=stl_its_col_break.font
    stl_its_col_break_font.name = 'Arial'
    stl_its_col_break_font.size = Pt(1)


    # Indentation Spacing
    pf_its_col_break = stl_its_col_break.paragraph_format

    pf_its_col_break.left_indent = Inches(0)
    pf_its_col_break.space_before=Pt(0)
    pf_its_col_break.space_after=Pt(0)

    pf_its_col_break.line_spacing_rule=WD_LINE_SPACING.EXACTLY



    ################## itssubhead ###############################################

    stlItssubHead = document.styles.add_style('itssubhead', WD_STYLE_TYPE.PARAGRAPH)
    stlItssubHead.quick_style =True
    # Font name

    stlItsSubHeadFont=stlItssubHead.font
    stlItsSubHeadFont.name = 'Arial'
    stlItsSubHeadFont.size = Pt(8)
    stlItsSubHeadFont.bold = False

    pfItsSubHead = stlItssubHead.paragraph_format

    # Indentation

    pfItsSubHead.left_indent = Inches(0.20)

    # Line spacing

    pfItsSubHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsSubHead.space_before=Pt(0)
    pfItsSubHead.space_after=Pt(0)
    pfItsSubHead.keep_together=True
    pfItsSubHead.keep_with_next=True
    # - - #


    ################## itsentry ###############################################

    stlitsentry= document.styles.add_style('itsentry', WD_STYLE_TYPE.PARAGRAPH)
    stlitsentry.quick_style =True
    # Font name

    stlitsentryFont=stlitsentry.font
    stlitsentryFont.name = 'Arial'
    stlitsentryFont.size = Pt(8)
    stlitsentryFont.bold = False


    pfstlitsentry = stlitsentry.paragraph_format

    # Indentation

    pfstlitsentry.first_line_indent = Inches(-0.2)
    pfstlitsentry.left_indent = Inches(0.6) # need 0.4 but indent from subhead seems to offset this one ?!?!

    # Line spacing

    pfstlitsentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfstlitsentry.keep_together=True
    pfstlitsentry.keep_with_next=True

    # - - #

    ################## docsymbol ###############################################

    stlDocSymbol = document.styles.add_style('docsymbol', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    stlDocSymbolFont=stlDocSymbol.font
    stlDocSymbolFont.name = 'Arial'
    stlDocSymbolFont.size = Pt(8)
    stlDocSymbolFont.bold = False


    pfdocsymbol = stlDocSymbol.paragraph_format

    # Indentation

    #pfdocsymbol.left_indent = Inches(0.40)

    # Line spacing

    #pfdocsymbol.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    # - - subroutines - - #

    def insert_its_head(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        
        #print(f"no of paragraphs is {len(document.paragraphs)}")
        p01=document.add_paragraph(itshead,style=stlItsHead)

    def insert_its_head_break(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        p011=document.add_paragraph(style=stl_its_col_break)
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        
    def insert_its_head_continued(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        p011=document.add_paragraph()
        p011.style=stl_its_col_break
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        p012=document.add_paragraph(itshead+" (continued)",style=stlItsHead)

    def insert_its_subhead(itssubhead):
            # Adding itssubhead content
            p10=document.add_paragraph(itssubhead,style=stlItssubHead)
            #p10.paragraph_format.space_after = Pt(0)

    def insert_its_entry(itsentry):
        
            #p2=document.add_paragraph(entry["entry"],style=stlitsentry)
            p20=document.add_paragraph(itsentry,style=stlitsentry)
            p20.add_run(" - ")
            #p20.paragraph_format.space_after = Pt(0)
            #try:

            
            #print(f"No of characters in entry for {entry['docsymbol']} is: {len(entry['entry'])}")
            #count1=1+((len(entry['docsymbol'])+len(entry['entry']))//52)
            #print(f"No of lines in entry for {entry['docsymbol']} is: {count1}")
            return p20
            
        #document.save('bgitpsubj'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    def insert_its_docsymbols(docsymbols, p30):
            #docsymbols=itsentry["docsymbols"]
            # retrieving size of the list
            #p30=document.add_paragraph(itsentry,style=stlitsentry)
            p30.paragraph_format.space_after = Pt(0)
            mySize=len(docsymbols)

            current=1

            for docsymbol in docsymbols:

                add_hyperlink(p30,docsymbol,Config.url_prefix+docsymbol)
                
                # Add the ';' and a space 
                if mySize>1 and current<mySize:
                    p30.add_run("; ")

                current+=1
            #p30.paragraph_format.space_after = Pt(0)



    # dxa Counters ##########################

    #def dxa_counter_symbol(docsymbols):
    #    chars=0
    #    for docsymbol in docsymbols:
    #        chars+=len(docsymbol)
    #   lines=1+(chars//52)
    #    #lines*fontsize*20+space after*20
    #    return lines*8*20+0*20

    def dxa_counter_itsentry(itsentry, docsymbols):
        chars=0
        for docsymbol in docsymbols:
            chars+=len(docsymbol)
        lines= 1+((len(itsentry)+chars)//45) 
        #lines*fontsize*20+space after*20
        return lines*8*20+0*20

    def dxa_counter_head(itshead):
        lines= 1+(len(itshead)//44) #37 is number of characters per line / on average ??
        #lines*fontsize*20+space after*20
        return lines*9*20+11*20 # 8 pts before and 3 after

    def dxa_counter_head_continued(itshead):
        lines = 1+((len(itshead)+len("(continued)"))//44)
        return lines*9*20+11*20

    def dxa_counter_subhead(itssubhead):
        lines=1+(len(itssubhead)//37)
        return lines*9*20+0*20
        
    #entries_counter=0
    dxa_counter=0
    dxa_page_length=10720# 8"*1440dxa 










    ################## WRITING THE DOCUMENT ###############################################

    # Adding the Header to the document

    p000=header.add_paragraph(myTitle.upper(), style='New Heading')
    p000.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding the sub Header to the document

    p001=header.add_paragraph(mySubTitle.upper(), style='New sub Heading')
    p001.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Breaks management
    paragraph_format = p001.paragraph_format
    paragraph_format.space_before = Pt(0)
    #paragraph_format.space_after = Pt(10)

    ##p003=header.add_paragraph("", style='New sub Heading')

    # Breaks management
    paragraph_format.keep_together = True
    paragraph_format.keep_with_next = True



    # data writing .. 

    myRecords=setOfData

    for record in myRecords:
        try:
            itshead=record["itshead"]
            subheading=record['subheading']
            itssubhead=subheading[0]["itssubhead"]
            itsentries=subheading[0]["itsentries"]
            try:
                itsentry=itsentries[0]["itsentry"]
            except:
                itsentry="symbol not used"
            docsymbols=itsentries[0]["docsymbols"]
            c_sh = dxa_counter_subhead(subheading[0]["itssubhead"])
            c_e=dxa_counter_itsentry(itsentry, docsymbols)
            c_h=dxa_counter_head(record)
            if dxa_counter+dxa_counter_head(record)+dxa_counter_subhead(subheading[0]["itssubhead"])+dxa_counter_itsentry(itsentry, docsymbols)>dxa_page_length:
                insert_its_head_break(record)
                dxa_counter=0
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(record)
            else:
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(record["itshead"])
                
            
            for mysubhead in subheading:
                            
                itssubhead=mysubhead["itssubhead"]
                itsentries=mysubhead["itsentries"]
                try:
                    itsentry=itsentries[0]["itsentry"]
                except:
                    itsentry="symbol not used"
                docsymbols=itsentries[0]["docsymbols"]
                if dxa_counter+dxa_counter_subhead(itssubhead)+dxa_counter_itsentry(itsentry, docsymbols)>dxa_page_length:
                    #insert_its_head_break(record)
                    dxa_counter=0
                    insert_its_head_continued(record)
                    dxa_counter+=dxa_counter_head_continued(record)
                    # Adding itssubhead content
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                else:
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                
                # array of arrays
                for itsentrie in itsentries :
                    #itsentrie is individual sub array 
                    # retrieve itsentry value
                    try:
                        itsentry=itsentrie["itsentry"]
                    except : 
                        itsentry="symbol not used"
                    try:
                        docsymbols=itsentrie["docsymbols"]
                    except : 
                        docsymbols="bl#"    
                    # Adding itssubhead content
                    if dxa_counter+dxa_counter_itsentry(itsentry, docsymbols)>dxa_page_length:
                        #insert_its_head_break(record)
                        dxa_counter=0
                        insert_its_head_continued(record)
                        dxa_counter+=dxa_counter_head_continued(record)

                        insert_its_subhead(itssubhead)
                        dxa_counter+=dxa_counter_subhead(itssubhead)
                        #if itsentry!="symbol not used":
                        p2=insert_its_entry(itsentry)
                        insert_its_docsymbols(docsymbols, p2)
                        dxa_counter+=dxa_counter_itsentry(itsentry, docsymbols)
                        '''
                            p20=document.add_paragraph(itsentry,style=stlitsentry)
                            p20.add_run(" - ")
                            
                            p20.paragraph_format.space_after = Pt(0)
                            mySize=len(docsymbols)

                            current=1

                            for docsymbol in docsymbols:

                                add_hyperlink1(p20,docsymbol,Config.url_prefix+docsymbol)
                                
                                # Add the ';' and a space 
                                if mySize>1 and current<mySize:
                                    p20.add_run("; ")

                                current+=1'''
                        # retrieve docsymbol value
                        
                        #dxa_counter+=dxa_counter_symbol(docsymbols)
                    else:
                        #if itsentry!="symbol not used":
                        p2=insert_its_entry(itsentry)
                        insert_its_docsymbols(docsymbols, p2)
                        dxa_counter+=dxa_counter_itsentry(itsentry, docsymbols)
                        '''
                            p20=document.add_paragraph(itsentry,style=stlitsentry)
                            p20.add_run(" - ")
                            mySize=len(docsymbols)

                            current=1

                            for docsymbol in docsymbols:

                                add_hyperlink1(p20,docsymbol,Config.url_prefix+docsymbol)
                                
                                # Add the ';' and a space 
                                if mySize>1 and current<mySize:
                                    p20.add_run("; ")

                                current+=1'''
                            
                        # retrieve docsymbol value
                        
                        #dxa_counter+=dxa_counter_symbol(docsymbols)
                    # Force a new line before next heading
            #p_new_line=document.add_paragraph(style=stlitsentry)
            #p_new_line.add_run("\n")
        except:
            pass           
        #document.save('bgitpitsc'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document

def generateWordDocITPITSP(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle=paramSubTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})

    # Creation of the word document
    document = Document()

    # Two columns display

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # Marging of the document

    section.top_margin = Inches(1)
    section.header_distance=Inches(0.5)
    section.bottom_margin = Inches(1)
    section.footer_distance=Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1) 

    ################## HEADER ###############################################

    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']

    # Font settings

    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    # Adding the header to the document

    header=document.sections[0].header

    ################## SUBHEADER ###############################################

    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_sub_heading_style.base_style = styles['Heading 1']

    # Font settings

    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    ################## itshead ###############################################

    stlItsHead = document.styles.add_style('itshead', WD_STYLE_TYPE.PARAGRAPH)
    stlItsHead.quick_style =True

    # Font name

    stlItsHeadFont=stlItsHead.font
    stlItsHeadFont.name = 'Arial'
    stlItsHeadFont.size = Pt(9)
    stlItsHeadFont.bold = True

    pfItsHead = stlItsHead.paragraph_format
    pfItsHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsHead.space_before=Pt(8)
    pfItsHead.space_after=Pt(3)
    pfItsHead.keep_together = True

    ################## column break style #####################################
    stl_its_col_break = document.styles.add_style('itsColBreak', WD_STYLE_TYPE.PARAGRAPH)
    stl_its_col_break.quick_style =True


    stl_its_col_break_font=stl_its_col_break.font
    stl_its_col_break_font.name = 'Arial'
    stl_its_col_break_font.size = Pt(1)


    # Indentation Spacing
    pf_its_col_break = stl_its_col_break.paragraph_format

    pf_its_col_break.left_indent = Inches(0)
    pf_its_col_break.space_before=Pt(0)
    pf_its_col_break.space_after=Pt(0)

    pf_its_col_break.line_spacing_rule=WD_LINE_SPACING.EXACTLY



    ################## itssubhead ###############################################

    stlItssubHead = document.styles.add_style('itssubhead', WD_STYLE_TYPE.PARAGRAPH)
    stlItssubHead.quick_style =True
    # Font name

    stlItsSubHeadFont=stlItssubHead.font
    stlItsSubHeadFont.name = 'Arial'
    stlItsSubHeadFont.size = Pt(8)
    stlItsSubHeadFont.bold = False

    pfItsSubHead = stlItssubHead.paragraph_format

    # Indentation

    pfItsSubHead.left_indent = Inches(0.20)

    # Line spacing

    pfItsSubHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsSubHead.space_before=Pt(0)
    pfItsSubHead.space_after=Pt(0)
    pfItsSubHead.keep_together=True
    pfItsSubHead.keep_with_next=True
    # - - #


    ################## itsentry ###############################################

    stlitsentry= document.styles.add_style('itsentry', WD_STYLE_TYPE.PARAGRAPH)
    stlitsentry.quick_style =True
    # Font name

    stlitsentryFont=stlitsentry.font
    stlitsentryFont.name = 'Arial'
    stlitsentryFont.size = Pt(8)
    stlitsentryFont.bold = False


    pfstlitsentry = stlitsentry.paragraph_format

    # Indentation

    pfstlitsentry.first_line_indent = Inches(-0.2)
    pfstlitsentry.left_indent = Inches(0.6) # need 0.4 but indent from subhead seems to offset this one ?!?!

    # Line spacing

    pfstlitsentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfstlitsentry.keep_together=True
    pfstlitsentry.keep_with_next=True

    # - - #

    ################## docsymbol ###############################################

    stlDocSymbol = document.styles.add_style('docsymbol', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    stlDocSymbolFont=stlDocSymbol.font
    stlDocSymbolFont.name = 'Arial'
    stlDocSymbolFont.size = Pt(8)
    stlDocSymbolFont.bold = False


    pfdocsymbol = stlDocSymbol.paragraph_format

    # Indentation

    #pfdocsymbol.left_indent = Inches(0.40)

    # Line spacing

    #pfdocsymbol.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    # - - subroutines - - #

    def insert_its_head(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        
        #print(f"no of paragraphs is {len(document.paragraphs)}")
        p01=document.add_paragraph(itshead,style=stlItsHead)

    def insert_its_head_break(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        p011=document.add_paragraph(style=stl_its_col_break)
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        
    def insert_its_head_continued(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        p011=document.add_paragraph()
        p011.style=stl_its_col_break
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        p012=document.add_paragraph(itshead+" (continued)",style=stlItsHead)

    def insert_its_subhead(itssubhead):
            # Adding itssubhead content
            p10=document.add_paragraph(itssubhead,style=stlItssubHead)
            #p10.paragraph_format.space_after = Pt(0)

    def insert_its_entry(itsentry):
        
            #p2=document.add_paragraph(entry["entry"],style=stlitsentry)
            p20=document.add_paragraph(itsentry,style=stlitsentry)
            #p20.paragraph_format.space_after = Pt(0)
            #try:

            
            #print(f"No of characters in entry for {entry['docsymbol']} is: {len(entry['entry'])}")
            #count1=1+((len(entry['docsymbol'])+len(entry['entry']))//52)
            #print(f"No of lines in entry for {entry['docsymbol']} is: {count1}")

            
        #document.save('bgitpsubj'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    def insert_its_docsymbols(docsymbols):
            #docsymbols=itsentry["docsymbols"]
            # retrieving size of the list
            p30=document.add_paragraph(itsentry,style=stlitsentry)
            p30.paragraph_format.space_after = Pt(0)
            mySize=len(docsymbols)

            current=1

            for docsymbol in docsymbols:

                add_hyperlink(p30,docsymbol,Config.url_prefix+docsymbol)
                
                # Add the ';' and a space 
                if mySize>1 and current<mySize:
                    p30.add_run("; ")

                current+=1
            #p30.paragraph_format.space_after = Pt(0)



    # dxa Counters ##########################

    def dxa_counter_symbol(docsymbols):
        chars=0
        for docsymbol in docsymbols:
            chars+=len(docsymbol)
        lines=1+(chars//52)
        #lines*fontsize*20+space after*20
        return lines*8*20+0*20

    def dxa_counter_itsentry(itsentry):
        lines= 1+(len(itsentry)//52) 
        #lines*fontsize*20+space after*20
        return lines*8*20+0*20

    def dxa_counter_head(itshead):
        lines= 1+(len(itshead)//48) #37 is number of characters per line / on average ??
        #lines*fontsize*20+space after*20
        return lines*9*20+11*20 # 8 pts before and 3 after

    def dxa_counter_head_continued(itshead):
        lines = 1+((len(itshead)+len("(continued)"))//48)
        return lines*9*20+11*20

    def dxa_counter_subhead(itssubhead):
        lines=1+(len(itssubhead)//42)
        return lines*9*20+0*20
        
    #entries_counter=0
    dxa_counter=0
    dxa_page_length=11120# 8"*1440dxa 

    ################## WRITING THE DOCUMENT ###############################################

    # Adding the Header to the document

    p000=header.add_paragraph(myTitle.upper(), style='New Heading')
    p000.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding the sub Header to the document

    p001=header.add_paragraph(mySubTitle.upper(), style='New sub Heading')
    p001.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Breaks management
    paragraph_format = p001.paragraph_format
    paragraph_format.space_before = Pt(0)
    #paragraph_format.space_after = Pt(10)

    ##p003=header.add_paragraph("", style='New sub Heading')

    # Breaks management
    paragraph_format.keep_together = True
    paragraph_format.keep_with_next = True



    # data writing .. 

    myRecords=setOfData

    for record in myRecords:
        try:
            itshead=record["itshead"]
            subheading=record['subheading']
            itssubhead=subheading[0]["itssubhead"]
            itsentries=subheading[0]["itsentries"]
            try:
                itsentry=itsentries[0]["itsentry"]
            except:
                itsentry=""
            docsymbols=itsentries[0]["docsymbols"]
            if dxa_counter+dxa_counter_head(itshead)+dxa_counter_subhead(subheading[0]["itssubhead"])+dxa_counter_itsentry(itsentry)+dxa_counter_symbol(docsymbols)>dxa_page_length:
                insert_its_head_break(record)
                dxa_counter=0
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(itshead)
            #document.save('bgitpitsp'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
            else:
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(itshead)
                
            
            for mysubhead in subheading:
                            
                itssubhead=mysubhead["itssubhead"]
                itsentries=mysubhead["itsentries"]
                try:
                    itsentry=itsentries[0]["itsentry"]
                except:
                    itsentry=""
                docsymbols=itsentries[0]["docsymbols"]
                if dxa_counter+dxa_counter_subhead(itssubhead)+dxa_counter_itsentry(itsentry)+dxa_counter_symbol(docsymbols)>dxa_page_length:
                    #insert_its_head_break(record)
                    dxa_counter=0
                    insert_its_head_continued(record)
                    dxa_counter+=dxa_counter_head_continued(record)
                    # Adding itssubhead content
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                else:
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                
                # array of arrays
                for itsentrie in itsentries :
                    #itsentrie is individual sub array 
                    # retrieve itsentry value
                    try:
                        itsentry=itsentrie["itsentry"]
                    except : 
                        itsentry=""
                    try:
                        docsymbols=itsentrie["docsymbols"]
                    except : 
                        docsymbols=" "    
                    # Adding itssubhead content
                    if dxa_counter+dxa_counter_itsentry(itsentry)+dxa_counter_symbol(docsymbols)>dxa_page_length:
                        #insert_its_head_break(record)
                        dxa_counter=0
                        insert_its_head_continued(record)
                        dxa_counter+=dxa_counter_head_continued(itshead)

                        insert_its_subhead(itssubhead)
                        dxa_counter+=dxa_counter_subhead(itssubhead)
                        if itsentry!="":
                            insert_its_entry(itsentry)
                            dxa_counter+=dxa_counter_itsentry(itsentry)
                        # retrieve docsymbol value
                        insert_its_docsymbols(docsymbols)
                        dxa_counter+=dxa_counter_symbol(docsymbols)
                    else:
                        if itsentry!="":
                            insert_its_entry(itsentry)
                            dxa_counter+=dxa_counter_itsentry(itsentry)
                        # retrieve docsymbol value
                        insert_its_docsymbols(docsymbols)
                        dxa_counter+=dxa_counter_symbol(docsymbols)
                    # Force a new line before next heading
            #p_new_line=document.add_paragraph(style=stlitsentry)
            #p_new_line.add_run("\n")
        except:
            pass           
            #document.save('bgitpsubj'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document    


def generateWordDocITPITSS(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):

    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle=paramSubTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})

    # Creation of the word document
    document = Document()

    # Two columns display

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # Marging of the document

    section.top_margin = Inches(1)
    section.header_distance=Inches(1)
    section.bottom_margin = Inches(1)
    section.footer_distance=Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1) 

    ################## HEADER ###############################################

    styles = document.styles
    stl_header = styles.add_style('Header1', WD_STYLE_TYPE.PARAGRAPH)
    #new_heading_style.base_style = styles['Heading 1']
    pf_header=stl_header.paragraph_format
    pf_header.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pf_header.space_before=Pt(0)
    pf_header.space_after=Pt(0)
    pf_header.keep_together = True

    # Font settings

    font = stl_header.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    # Adding the header to the document

    header=document.sections[0].header

    ################## SUBHEADER ###############################################

    #new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    #new_sub_heading_style.base_style = styles['Heading 1']

    # Font settings

    #font = new_sub_heading_style.font
    #font.name = 'Arial'
    #font.size = Pt(8)
    #font.bold = False
    #font.color.rgb = RGBColor(0, 0, 0)

    ################## itshead ###############################################

    stlItsHead = document.styles.add_style('itshead', WD_STYLE_TYPE.PARAGRAPH)
    stlItsHead.quick_style =True

    # Font name

    stlItsHeadFont=stlItsHead.font
    stlItsHeadFont.name = 'Arial'
    stlItsHeadFont.size = Pt(9)
    stlItsHeadFont.bold = True

    pfItsHead = stlItsHead.paragraph_format
    pfItsHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsHead.space_before=Pt(8)
    pfItsHead.space_after=Pt(3)
    pfItsHead.keep_together = True

    ################## column break style #####################################
    stl_its_col_break = document.styles.add_style('itsColBreak', WD_STYLE_TYPE.PARAGRAPH)
    stl_its_col_break.quick_style =True


    stl_its_col_break_font=stl_its_col_break.font
    stl_its_col_break_font.name = 'Arial'
    stl_its_col_break_font.size = Pt(1)


    # Indentation Spacing
    pf_its_col_break = stl_its_col_break.paragraph_format

    pf_its_col_break.left_indent = Inches(0)
    pf_its_col_break.space_before=Pt(0)
    pf_its_col_break.space_after=Pt(0)

    pf_its_col_break.line_spacing_rule=WD_LINE_SPACING.EXACTLY



    ################## itssubhead ###############################################

    stlItssubHead = document.styles.add_style('itssubhead', WD_STYLE_TYPE.PARAGRAPH)
    stlItssubHead.quick_style =True
    # Font name

    stlItsSubHeadFont=stlItssubHead.font
    stlItsSubHeadFont.name = 'Arial'
    stlItsSubHeadFont.size = Pt(8)
    stlItsSubHeadFont.bold = False

    pfItsSubHead = stlItssubHead.paragraph_format

    # Indentation

    pfItsSubHead.left_indent = Inches(0.20)

    # Line spacing

    pfItsSubHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsSubHead.space_before=Pt(0)
    pfItsSubHead.space_after=Pt(0)
    pfItsSubHead.keep_together=True
    pfItsSubHead.keep_with_next=True
    # - - #


    ################## itsentry ###############################################

    stlitsentry= document.styles.add_style('itsentry', WD_STYLE_TYPE.PARAGRAPH)
    stlitsentry.quick_style =True
    # Font name

    stlitsentryFont=stlitsentry.font
    stlitsentryFont.name = 'Arial'
    stlitsentryFont.size = Pt(8)
    stlitsentryFont.bold = False


    pfstlitsentry = stlitsentry.paragraph_format

    # Indentation

    pfstlitsentry.first_line_indent = Inches(-0.2)
    pfstlitsentry.left_indent = Inches(0.6) # need 0.4 but indent from subhead seems to offset this one ?!?!

    # Line spacing

    pfstlitsentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfstlitsentry.keep_together=True
    pfstlitsentry.keep_with_next=True

    # - - #

    ################## docsymbol ###############################################

    stlDocSymbol = document.styles.add_style('docsymbol', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    stlDocSymbolFont=stlDocSymbol.font
    stlDocSymbolFont.name = 'Arial'
    stlDocSymbolFont.size = Pt(8)
    stlDocSymbolFont.bold = False


    pfdocsymbol = stlDocSymbol.paragraph_format

    # Indentation

    #pfdocsymbol.left_indent = Inches(0.40)

    # Line spacing

    #pfdocsymbol.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    # - - subroutines - - #

    def insert_its_head(record, dxa_counter):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        
        #print(f"no of paragraphs is {len(document.paragraphs)}")
        p01=document.add_paragraph(itshead,style=stlItsHead)
        if dxa_counter<=180:
            p01.paragraph_format.space_before=Pt(0)

    def insert_its_head_break(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        p011=document.add_paragraph(style=stl_its_col_break)
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        
    def insert_its_head_continued(record):
        try :
            itshead= record['itshead']
        except :
            itshead=""
        
        # Adding itshead content
        p011=document.add_paragraph()
        p011.style=stl_its_col_break
        column_break_run = p011.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        p012=document.add_paragraph(itshead+" (continued)",style=stlItsHead)
        p012.paragraph_format.space_before=Pt(0)

    def insert_its_subhead(itssubhead):
            # Adding itssubhead content
            p10=document.add_paragraph(itssubhead,style=stlItssubHead)
            #p10.paragraph_format.space_after = Pt(0)

    def insert_its_entry(itsentry):
        
            #p2=document.add_paragraph(entry["entry"],style=stlitsentry)
            p20=document.add_paragraph(itsentry,style=stlitsentry)
            p20.add_run(" - ")
            #p20.paragraph_format.space_after = Pt(0)
            #try:

            
            #print(f"No of characters in entry for {entry['docsymbol']} is: {len(entry['entry'])}")
            #count1=1+((len(entry['docsymbol'])+len(entry['entry']))//52)
            #print(f"No of lines in entry for {entry['docsymbol']} is: {count1}")
            return p20
            
        #document.save('bgitpsubj'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    def insert_its_docsymbols(docsymbols, p30):
            #docsymbols=itsentry["docsymbols"]
            # retrieving size of the list
            #p30=document.add_paragraph(itsentry,style=stlitsentry)
            p30.paragraph_format.space_after = Pt(0)
            mySize=len(docsymbols)

            current=1

            for docsymbol in docsymbols:

                add_hyperlink(p30,docsymbol,Config.url_prefix+docsymbol)
                
                # Add the ';' and a space 
                if mySize>1 and current<mySize:
                    p30.add_run("; ")

                current+=1
            #p30.paragraph_format.space_after = Pt(0)



    # dxa Counters ##########################

    #def dxa_counter_symbol(docsymbols):
    #    chars=0
    #    for docsymbol in docsymbols:
    #        chars+=len(docsymbol)
    #   lines=1+(chars//52)
    #    #lines*fontsize*20+space after*20
    #    return lines*8*20+0*20



    def dxa_counter_head(itshead, dxa_counter):
        lines= 1+(len(itshead)//38) #37 is number of characters per line / on average ??
        #lines*fontsize*20+space after*20
        if dxa_counter < 8000:
            return lines*9*20+11*20 # 8 pts before and 3 after
        else:
            return lines*9*20+3*20 # 8 pts before and 3 after

    def dxa_counter_head_continued(itshead):
        lines = 1+((len(itshead)+len("(continued)"))//38)
        return lines*9*20+3*20

    def dxa_counter_subhead(itssubhead):
        lines=1+(len(itssubhead)//50)
        return lines*8*20+0*20

    def dxa_counter_itsentries(itsentries):
        dxa_temp=0
        for itsentry in itsentries:
            dxa_temp+=dxa_counter_itsentry(itsentry["itsentry"],itsentry["docsymbols"])
        return dxa_temp
        
    def dxa_counter_itsentry(itsentry, docsymbols):
        chars=0
        for i,docsymbol in enumerate(docsymbols):
            chars+=len(docsymbol)
        lines= 1+(len(itsentry)+(((i+1)*3)+chars))//50
        #print(f"doc symbol len is {chars}")
        #print(f"for {itsentry} number of characters is {len(itsentry)+chars} and calc lines is {1+((len(itsentry)+chars)//43) }")
        #print(f"number of lines is {lines}")
        #lines*fontsize*20+space after*20
        return lines*8*20+0*20
    def dxa_counter_docsymbols(docsymbols):
        chars=0
        for i,docsymbol in enumerate(docsymbols):
            chars+=len(docsymbol)
        lines= 1+((((i+1)*3)+chars))//50
        #print(f"doc symbol len is {chars}")
        #print(f"for {itsentry} number of characters is {len(itsentry)+chars} and calc lines is {1+((len(itsentry)+chars)//43) }")
        #print(f"number of lines is {lines}")
        #lines*fontsize*20+space after*20
        return lines*8*20+0*20


    #entries_counter=0
    dxa_counter=0
    dxa_page_length=9920# 7"*1440dxa - 5 linija *180 




    ################## WRITING THE DOCUMENT ###############################################

    # Adding the Header to the document

    p000=header.paragraphs[0]
    p000.style='Header1'
    p000.text=myTitle.upper()
    p000.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding the sub Header to the document
    run=p000.add_run()
    run.add_break()
    p000.add_run(mySubTitle.upper())
    #p001.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run=p000.add_run()
    run.add_break()
    run.add_break()
    #p001.add_run("\n")

    # Breaks management
    #paragraph_format = p001.paragraph_format
    #paragraph_format.space_before = Pt(0)
    #paragraph_format.space_after = Pt(10)

    ##p003=header.add_paragraph("", style='New sub Heading')

    # Breaks management
    #paragraph_format.keep_together = True
    #paragraph_format.keep_with_next = True



    # data writing .. 

    myRecords=setOfData
    lines_on_page=0
    for record in myRecords:
        try:
            itshead=record["itshead"]
            subheading=record['subheading']
            itssubhead=subheading[0]["itssubhead"]
            itsentries=subheading[0]["itsentries"]
            try:
                itsentry=itsentries[0]["itsentry"]
            except:
                itsentry="symbol not used"
            docsymbols=itsentries[0]["docsymbols"]
            c_sh = dxa_counter_subhead(subheading[0]["itssubhead"])
            c_ie=dxa_counter_itsentries(itsentries)
            c_h=dxa_counter_head(record, dxa_counter)
            #document.save('bgitpitss'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
            if dxa_counter+dxa_counter_head(itshead,dxa_counter)+dxa_counter_subhead(subheading[0]["itssubhead"])+dxa_counter_itsentries(itsentries)>dxa_page_length:
                insert_its_head_break(record)
                dxa_counter=0
                lines_on_page=0
                insert_its_head(record, dxa_counter)# to remove 8 points before, when its head is at the start of the column
                dxa_counter+=dxa_counter_head(itshead, dxa_counter)
            else:
                insert_its_head(record, dxa_counter)
                dxa_counter+=dxa_counter_head(itshead, dxa_counter)
                
            
            for mysubhead in subheading:
                            
                itssubhead=mysubhead["itssubhead"]
                itsentries=mysubhead["itsentries"]
                try:
                    itsentry=itsentries[0]["itsentry"]
                except:
                    itsentry="symbol not used"
                docsymbols=itsentries[0]["docsymbols"]
                #print(f" number of lines for {itssubhead} is {(dxa_counter_subhead(itssubhead)+dxa_counter_itsentries(itsentries))/160}")
                lines_on_page+=(dxa_counter_subhead(itssubhead)+dxa_counter_itsentries(itsentries))/160
                #print(f"total no of lines is {lines_on_page}")
                if dxa_counter+dxa_counter_subhead(itssubhead)+dxa_counter_itsentries(itsentries)>dxa_page_length:
                    #insert_its_head_break(record)
                    dxa_counter=0
                    lines_on_page=0
                    insert_its_head_continued(record)
                    dxa_counter+=dxa_counter_head_continued(itshead)
                    # Adding itssubhead content
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                else:
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)
                
                # array of arrays
                for itsentrie in itsentries :
                    #itsentrie is individual sub array 
                    # retrieve itsentry value
                    try:
                        itsentry=itsentrie["itsentry"]
                    except : 
                        itsentry="symbol not used"
                    try:
                        docsymbols=itsentrie["docsymbols"]
                    except : 
                        docsymbols=""    
                    # Adding itssubhead content
                    if dxa_counter+dxa_counter_itsentry(itsentry, docsymbols)>dxa_page_length:
                        #insert_its_head_break(record)
                        dxa_counter=0
                        lines_on_page=0
                        insert_its_head_continued(record)
                        dxa_counter+=dxa_counter_head_continued(itshead)

                        insert_its_subhead(itssubhead)
                        dxa_counter+=dxa_counter_subhead(itssubhead)
                        #if itsentry!="symbol not used":
                        p2=insert_its_entry(itsentry)
                        insert_its_docsymbols(docsymbols, p2)
                        dxa_counter+=dxa_counter_itsentry(itsentry, docsymbols)
                        '''
                            p20=document.add_paragraph(itsentry,style=stlitsentry)
                            p20.add_run(" - ")
                            
                            p20.paragraph_format.space_after = Pt(0)
                            mySize=len(docsymbols)

                            current=1

                            for docsymbol in docsymbols:

                                add_hyperlink1(p20,docsymbol,Config.url_prefix+docsymbol)
                                
                                # Add the ';' and a space 
                                if mySize>1 and current<mySize:
                                    p20.add_run("; ")

                                current+=1'''
                        # retrieve docsymbol value
                        
                        #dxa_counter+=dxa_counter_symbol(docsymbols)
                    else:
                        #if itsentry!="symbol not used":
                        p2=insert_its_entry(itsentry)
                        insert_its_docsymbols(docsymbols, p2)
                        dxa_counter+=dxa_counter_itsentry(itsentry, docsymbols)
                        '''
                            p20=document.add_paragraph(itsentry,style=stlitsentry)
                            p20.add_run(" - ")
                            mySize=len(docsymbols)

                            current=1

                            for docsymbol in docsymbols:

                                add_hyperlink1(p20,docsymbol,Config.url_prefix+docsymbol)
                                
                                # Add the ';' and a space 
                                if mySize>1 and current<mySize:
                                    p20.add_run("; ")

                                current+=1'''
                            
                        # retrieve docsymbol value
                        
                        #dxa_counter+=dxa_counter_symbol(docsymbols)
                    # Force a new line before next heading
            #p_new_line=document.add_paragraph(style=stlitsentry)
            #p_new_line.add_run("\n")
        except:
            pass           
        #document.save('bgitpitsc'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document  

def generateWordDocITPRES(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Function to keep the header visible for the table

    def set_repeat_table_header(row):

        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle1=paramSubTitle
    #setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection}).sort("sortkey1",1)
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection}).sort("sortkey1",1).collation(Collation(locale='en',numericOrdering=True))
    
    # Creation of the word document
    document = Document()
    
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Adding the header to the document
    
    header=document.sections[0].header
    
    ################## SUBHEADER ###############################################
    
    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    #new_sub_heading_style.base_style = styles['Heading 1']
    
    
    # Font settings
    
    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    ################## FIRST LINE ###############################################
    
    first_line = styles.add_style('first line', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font settings
    
    font = first_line.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.italic=True
    font.color.rgb = RGBColor(0, 0, 0)
    
    ################## SPECIAL ###############################################
    
    special = styles.add_style('special', WD_STYLE_TYPE.PARAGRAPH)
    pf = special.paragraph_format
    pf.left_indent = Inches(0.40)
    # Font settings
    font = special.font
    font.underline=True
    
    
    ################## TABLE CONTENT ###############################################
    
    tableContent_style = styles.add_style('tableContent', WD_STYLE_TYPE.PARAGRAPH)
    tableContent_style.base_style = styles['Normal']
    
    # Font settings
    
    font = tableContent_style.font
    font.name = 'Arial'
    font.size = Pt(8)

    
    # Hanging indent    

    pf = tableContent_style.paragraph_format
    #pf.first_line_indent = Inches(-0.02)
    
    pf.first_line_indent = Inches(-0.06)

    ################## TABLE CONTENT 1 ###############################################
    
    tableContent_style1 = styles.add_style('tableContent1', WD_STYLE_TYPE.PARAGRAPH)
    tableContent_style1.base_style = styles['Normal']
    
    # Font settings
    
    font = tableContent_style1.font
    font.name = 'Arial'
    font.size = Pt(8)

    
    # Hanging indent    

    pf1 = tableContent_style1.paragraph_format
    pf1.left_indent = Pt(10)
    pf1.space_before = Pt(1)
    
    #pf.first_line_indent = Inches(10)

    ################## WRITING THE DOCUMENT ###############################################
    
    # Header Generation
    
    if (bodysession[0] in ["A","E"]):
        #p=header.add_paragraph(mySubTitle1.upper(),style='New Heading')
        p=header.add_paragraph(myTitle.upper(),style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    else : 
        p=header.add_paragraph(myTitle.upper(), style='New Heading')
        #p=header.add_paragraph(myTitle, style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        # Adding the sub Header to the document
    
        #p=header.add_paragraph(mySubTitle1, style='New sub Heading')
        #p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
 
    # First line generation

    paragraph_format = p.paragraph_format
    paragraph_format.space_after = Pt(20)
    #paragraph_format.left_indent=Cm(0.508)

    p=document.add_paragraph("Vote reads Yes-No-Abstain",style='first line')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = p.paragraph_format
    #paragraph_format.left_indent=Inches(0.35)

    # Definition of the column names
    myRecords=setOfData

    if (bodysession[0] in ["A","E"]):

        # creation of the table
        table = document.add_table(rows=1, cols=5)

        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        #table.allow_autofit=True

        # set the header visible
        set_repeat_table_header(table.rows[0])

        # Retrieve the first line
        hdr_cells = table.rows[0].cells

        firstTitle= myRecords[0]["docsymbol"].split("/")
        baseUrl1=firstTitle[0] + "/" + firstTitle[1] + "/" + firstTitle[2] + "/" 
        firstlineValue= myRecords[0]["meeting"].split(".")
        baseUrl2="("+firstlineValue[0]+".-)"
        
        run = hdr_cells[0].paragraphs[0].add_run(baseUrl1)
        run.underline=True
        
        run = hdr_cells[1].paragraphs[0].add_run('Title')
        tc = hdr_cells[1]._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        start = OxmlElement('w:start')
        
        tcMar.set(qn('w:w'), "0")
        tcMar.set(qn('w:type'),"dxa")
              
        tcMar.append(start)
        tcPr.append(tcMar)
        run.underline=True     
        
        hdr_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        #run = hdr_cells[2].paragraphs[0].add_run('Meeting / Date' +  "               " +baseUrl2)
        run = hdr_cells[2].paragraphs[0].add_run('Meeting / Date')
        run.underline=True
        run1=hdr_cells[2].paragraphs[0].add_run("               " +baseUrl2)
        
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[3].paragraphs[0].add_run('A.I. No.')
        run.underline=True
        
        hdr_cells[4].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[4].paragraphs[0].add_run('Vote')
        run.underline=True
        

        # writing the others lines

        for record in myRecords:
            firstTitle=""
            firstlineValue=""
            firstTitle= record["docsymbol"].split("/")

            firstlineValue= record["meeting"].split(".")
            row_cells = table.add_row().cells

            add_hyperlink(row_cells[0].paragraphs[0],firstTitle[3],Config.url_prefix+firstTitle[0]+"/"+firstTitle[1]+"/"+firstTitle[2]+"/"+firstTitle[3])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text=record["title"]
            row_cells[1].paragraphs[0].style="tableContent" 

            ### based on Jennifer's update on 12 May 2021
            if ''.join(firstlineValue) !="":
                add_hyperlink(row_cells[2].paragraphs[0],firstlineValue[1],Config.url_prefix+firstlineValue[0]+"."+firstlineValue[1])
                row_cells[2].paragraphs[0].add_run( " / " + record["votedate"])
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row_cells[2].paragraphs[0].add_run(record["votedate"])
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ############################################
            
            row_cells[3].text=record["ainumber"]
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ############################################
            
            row_cells[4].text=record["vote"]
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Definition of the size of the column
        #  #widths = (Inches(0.73), Inches(3.30), Inches(1.05), Inches(0.61), Inches(1.10)) second version
        widths = (Inches(0.43), Inches(5.00), Inches(1.25), Inches(1.0), Inches(1.10))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    else :

        # creation of the table
        table = document.add_table(rows=1, cols=4)

        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        #table.allow_autofit=True

        # set the header visible
        set_repeat_table_header(table.rows[0])

        # Retrieve the first line
        hdr_cells = table.rows[0].cells

        firstTitle= myRecords[0]["docsymbol"].split("/")
        baseUrl1=firstTitle[0] + "/" + firstTitle[1] + "/"
    
        firstlineValue= (myRecords[0]["meeting"].split("."))[0].split("/")
        if ''.join(firstlineValue) !="":
            baseUrl2="  ("+firstlineValue[0]+"/"+ firstlineValue[2] +".-)"

        hdr_cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[0].paragraphs[0].add_run(baseUrl1)
        run.underline=True
        
        run = hdr_cells[1].paragraphs[0].add_run('Subject')
        tc2 = hdr_cells[1]._tc
        tcPr2 = tc2.get_or_add_tcPr()
        tcMar2 = OxmlElement('w:tcMar')
        start2 = OxmlElement('w:start')
        
        tcMar2.set(qn('w:w'), "0")
        tcMar2.set(qn('w:type'),"dxa")
              
        tcMar2.append(start2)
        tcPr2.append(tcMar2)
        run.underline=True


        hdr_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[2].paragraphs[0].add_run('Meeting / Date, '+myRecords[0]["voteyear"] )
        run.underline=True
        if ''.join(firstlineValue) !="":
            run1=hdr_cells[2].paragraphs[0].add_run(" "+baseUrl2)

        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[3].paragraphs[0].add_run('Vote')
        run.underline=True
       

        # writing the others lines

        for record in myRecords:
            firstTitle=""
            firstlineValue=""
            firstTitle= record["docsymbol"].split("/")
            firstlineValue= record["meeting"].split(".")
            firstlineValuePlus=firstlineValue[0].split("/")
            row_cells = table.add_row().cells

            add_hyperlink(row_cells[0].paragraphs[0],firstTitle[2],Config.url_prefix+firstTitle[0]+"/"+firstTitle[1]+"/"+firstTitle[2]) 
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].text=record["subject"].upper()
            
            ########
            tc2 = row_cells[1]._tc
            tcPr2 = tc2.get_or_add_tcPr()
            tcMar2 = OxmlElement('w:tcMar')
            start2 = OxmlElement('w:start')
            
            tcMar2.set(qn('w:w'), "0")
            tcMar2.set(qn('w:type'),"dxa")
                
            tcMar2.append(start2)
            tcPr2.append(tcMar2)

            ########
            
            paragraph_format = row_cells[1].paragraphs[0].paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            row_cells[1].add_paragraph(record["subjectsubtitle"],style="tableContent1")
            
            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            if ''.join(firstlineValue) !="":
                add_hyperlink(row_cells[2].paragraphs[0],firstlineValue[1],Config.url_prefix+firstlineValuePlus[0]+"/"+ firstlineValuePlus[2]+ "."+firstlineValue[1])
                row_cells[2].paragraphs[0].add_run( " / " + record["votedate"])
            else:
                row_cells[2].paragraphs[0].add_run(record["votedate"])
            row_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            #row_cells[3].text=record["vote"]
            row_cells[3].paragraphs[0].text=record["vote"]
            


        # Definition of the size of the column

        widths = (Inches(1.3), Inches(5), Inches(1.7),Inches(0.75))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    # Definition of the font

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                # Adding styling
                for run in paragraph.runs:
                    font = run.font
                    font.name="Arial"
                    font.size= Pt(8)

    # Save the document
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document  

def generateWordDocITPSUBJ(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle=paramSubTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})

    # Creation of the word document
    document = Document()

    # Two columns display

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # Marging of the document

    section.top_margin = Inches(1)
    section.header_distance=Inches(0.5)
    section.bottom_margin = Inches(1)
    section.footer_distance=Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1) 



    ################## HEADER ###############################################

    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']


    # Font settings

    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    # Adding the header to the document

    header=document.sections[0].header

    pfnew_heading_style = new_heading_style.paragraph_format
    pfnew_heading_style.line_spacing_rule =  WD_LINE_SPACING.SINGLE
   
    ################## SUBHEADER ############################################### 

    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_sub_heading_style.base_style = styles['Heading 1']
    # Font settings

    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

    ################## itshead ############################################### these are subjects

    stlItsHead = document.styles.add_style('itshead', WD_STYLE_TYPE.PARAGRAPH)
    stlItsHead.quick_style =True

    # Font name

    stlItsHeadFont=stlItsHead.font
    stlItsHeadFont.name = 'Arial'
    stlItsHeadFont.size = Pt(9)
    stlItsHeadFont.bold = True

    pfItsHead = stlItsHead.paragraph_format
    pfItsHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsHead.space_before=Pt(8)
    pfItsHead.space_after=Pt(3)
    pfItsHead.keep_together = True


    # Line spacing

    ################## column break style #####################################
    stl_its_col_break = document.styles.add_style('itsColBreak', WD_STYLE_TYPE.PARAGRAPH)
    stl_its_col_break.quick_style =True


    stl_its_col_break_font=stl_its_col_break.font
    stl_its_col_break_font.name = 'Arial'
    stl_its_col_break_font.size = Pt(1)


    # Indentation Spacing
    pf_its_col_break = stl_its_col_break.paragraph_format

    pf_its_col_break.left_indent = Inches(0)
    pf_its_col_break.space_before=Pt(0)
    pf_its_col_break.space_after=Pt(0)

    pf_its_col_break.line_spacing_rule=WD_LINE_SPACING.EXACTLY


    ################## itssubhead ############################################### this is for grouping

    stlItssubHead = document.styles.add_style('itssubhead', WD_STYLE_TYPE.PARAGRAPH)
    stlItssubHead.quick_style =True

    # Font name

    stlItsSubHeadFont=stlItssubHead.font
    stlItsSubHeadFont.name = 'Arial'
    stlItsSubHeadFont.size = Pt(9)
    stlItsSubHeadFont.bold = True
    stlItsSubHeadFont.underline = True

    pfItsSubHead = stlItssubHead.paragraph_format
    pfItsSubHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfItsSubHead.keep_together=True
    pfItsSubHead.space_before=Pt(3)
    pfItsSubHead.space_after=Pt(0)

    # Indentation

    pfItsSubHead.left_indent = Inches(0)

    ################## itsentry ############################################### entries styling

    stlitsentry= document.styles.add_style('itsentry', WD_STYLE_TYPE.PARAGRAPH)
    stlitsentry.quick_style =True
    # Font name

    stlitsentryFont=stlitsentry.font
    stlitsentryFont.name = 'Arial'
    stlitsentryFont.size = Pt(8)
    stlitsentryFont.bold = False
    pfstlitsentry = stlitsentry.paragraph_format

    # Indentation

    pfstlitsentry.left_indent = Inches(0.2)
    pfstlitsentry.first_line_indent  = Inches(-0.1)

    # Line spacing

    pfstlitsentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfstlitsentry.keep_together=True
    pfstlitsentry.space_before=Pt(5)
    pfstlitsentry.space_after=Pt(0)

    ################## note ############################################### notes styling

    stlNote = document.styles.add_style('note', WD_STYLE_TYPE.PARAGRAPH)
    stlNote.quick_style =True

    # Font name

    stlNoteFont=stlNote.font
    stlNoteFont.name = 'Arial'
    stlNoteFont.size = Pt(8)

    pfNote = stlNote.paragraph_format

    # Indentation

    pfNote.left_indent = Inches(0.1)
    pfNote.first_line_indent = Inches(0.3)

    # Line spacing

    pfNote.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    pfNote.keep_together=True
    pfNote.space_before=Pt(0)
    pfNote.space_after=Pt(0)


    ################## WRITING THE DOCUMENT ############################################### 

    # Adding the Header to the document

    p=header.add_paragraph(myTitle.upper(), style='New Heading')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding the sub Header to the document

    p1=header.add_paragraph(mySubTitle.upper(), style='New sub Heading')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Breaks management
    paragraph_format = p1.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(10)
    paragraph_format.keep_together = True
    paragraph_format.keep_with_next = True

    myRecords=setOfData

    def insert_its_head(record):
        try :
            itshead= record['head']
        except :
            itshead=""
        
        # Adding itshead content
        
        p=document.add_paragraph(itshead,style=stlItsHead)

    def insert_its_head_break(record):
        try :
            itshead= record['head']
        except :
            itshead=""
        
        p0=document.add_paragraph()
        p0.style=stl_its_col_break
        column_break_run = p0.add_run()
        column_break_run.add_break(WD_BREAK.COLUMN)
        
    def insert_its_head_continued(record):
        try :
            itshead= record['head']
        except :
            itshead=""
        
        # Adding itshead content
        p=document.add_paragraph(itshead+" (continued)",style=stlItsHead)

    def insert_its_subhead(itssubhead):
            # Adding itssubhead content
            p1=document.add_paragraph(itssubhead,style=stlItssubHead)

    def insert_entry_note(entry):
        
        if entry["note"]!="":
            #p2=document.add_paragraph(entry["entry"],style=stlitsentry)
            p2=document.add_paragraph(style=stlitsentry)
            #dual document symbol
            myEntry=entry["docsymbol"].split(" (S/")
            if len(myEntry)==2:
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (S/"+myEntry[1])
            elif len(entry["docsymbol"].split(" (E/"))==2:
                myEntry=entry["docsymbol"].split(" (E/")
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (E/"+myEntry[1])
            elif len(entry["docsymbol"].split(" (A/"))==2:
                myEntry=entry["docsymbol"].split(" (A/")
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (A/"+myEntry[1])
            else:
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])


            p2.add_run(" ")
            p2.add_run(entry["entry"])
            #print(f"No of characters in entry for {entry['docsymbol']} is: {len(entry['entry'])}") 
            count1=1+((len(entry['docsymbol'])+len(entry['entry']))//52)
            #print(f"No of lines in entry for {entry['docsymbol']} is: {count1}")
            #Breaks management
            paragraph_format = p2.paragraph_format
            paragraph_format.keep_with_next = True
        
        

            insert_note(entry)
            

        else:
            #p2=document.add_paragraph(entry["entry"],style=stlitsentry)
            p2=document.add_paragraph(style=stlitsentry)
            #try:

            #dual document symbol
            myEntry=entry["docsymbol"].split(" (S/")
            if len(myEntry)==2:
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (S/"+myEntry[1])
            elif len(entry["docsymbol"].split(" (E/"))==2:
                myEntry=entry["docsymbol"].split(" (E/")
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (E/"+myEntry[1])
            elif len(entry["docsymbol"].split(" (A/"))==2:
                myEntry=entry["docsymbol"].split(" (A/")
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])
                p2.add_run(" (A/"+myEntry[1])
            else:
                add_hyperlink1(p2,myEntry[0],Config.url_prefix+myEntry[0])

        
            p2.add_run(" ")
            my_string_with_only_one_period = get_note_with_one_periode(entry["entry"])
            p2.add_run(my_string_with_only_one_period)
            #print(f"No of characters in entry for {entry['docsymbol']} is: {len(entry['entry'])}")
            count1=1+((len(entry['docsymbol'])+len(my_string_with_only_one_period))//52)
            #print(f"No of lines in entry for {entry['docsymbol']} is: {count1}")

    def insert_note(entry):
        p3=document.add_paragraph(entry["note"],style=stlNote)
        #print(f"No of characters in note for {entry['docsymbol']} is: {len(entry['note'])}")
        count2=1+(len(entry['note'])//52)
        #print(f"No of lines in note for {entry['docsymbol']} is: {count2}")
        #Breaks management
        paragraph_format = p3.paragraph_format
        paragraph_format.keep_together = True

    def get_note_with_one_periode(myString):
        my_lenght=len(myString)
        if (myString[my_lenght-1]==myString[my_lenght-2]) and (myString[my_lenght-1]=="."):
            return myString[:-1]
        return myString    
            

    # to break columns .. 
    nextPage={}

    def dxa_counter_entry_note(entry):
        if entry["note"]!="":
            count_n=1+(len(entry['note'])//48)
            count_e=1+((len(entry['docsymbol'])+len(entry['entry']))//48)
            lines=count_n+count_e
        else:
            count_n=0
            count_e=1+((len(entry['docsymbol'])+len(entry['entry']))//48)
            lines=count_n+count_e
            #lines*fontsize*20+space after*20
        return lines*8*20+5*20

    def dxa_counter_head(head):
        lines= 1+(len(head)//37) #37 is number of characters per line / on average ??
        #lines*fontsize*20+space after*20
        return lines*9*20+11*20 # 8 pts before and 3 after

    def dxa_counter_head_continued(head):
        lines = 1+((len(head)+len("(continued)"))//37)
        return lines*9*20+11*20

    def dxa_counter_subhead(subhead):
        lines=1+(len(subhead)//37)
        return lines*9*20+3*20
        
    entries_counter=0
    dxa_counter=0
    dxa_page_length=11111# 8"*1440dxa - 300

    # this function will check if the new entry is already 
    # inside the doc

    def formatEntry(entries):

         # only if we have to deal with S
        if (bodysession[0]=="S"):

            listWithoutNote=[]
            listWithoutNote1=[]
            listEntryUpdate=entries
            
            # retrieve records without note
            for entry in entries:
                if (entry["note"]=="" ):     
                    listWithoutNote.append(entry)

            for entry in entries:        
                if (entry["note"]==""):      
                    listWithoutNote1.append(entry)
                
                    
            # first user case 
            for withoutNote in listWithoutNote:
                for entry in listEntryUpdate:
                    if withoutNote["docsymbol"]==entry["docsymbol"] and withoutNote["entry"]==entry["entry"] and withoutNote["note"]!=entry["note"] :
                        # remove this note from the initial list
                        entries.remove(withoutNote)
            

            # second user case 
            for WithoutNote1 in listWithoutNote1:
                occurence=0
                for entry in listEntryUpdate:
                    if WithoutNote1["docsymbol"]==entry["docsymbol"] and WithoutNote1["entry"]==entry["entry"] and entry["note"]=="" :
                        occurence=occurence+1
                        if occurence>1 :
                            entries.remove(entry)

    # subheader List
    # myDocSymbolList=[]
    
    def check_presence_value(my_array,my_value):
        return_value=False
        for value in my_array:
            if (value["docsymbol"]==my_value):
                return_value=True
        return return_value
        
    
    def sort_entry(my_array_of_entries):

        array_Add=[] # Array with Add.XXXX 
        array_Corr=[] # Array with Corr.XXX
        array_Final=[] # Final Array with Add and Corr
        array_already_processed=[]
        
        
        # retrieving Add. and copy inside array_Add
        for value_add in my_array_of_entries:
            if (value_add["docsymbol"].find("Add.")!= -1):
                array_Add.append(value_add)
                
        # retrieving Add. and copy inside array_Add
        for value_corr in my_array_of_entries:
            if (value_corr["docsymbol"].find("Corr.")!= -1):
                array_Corr.append(value_corr)
                
        # building final array
        for value in my_array_of_entries:

            if value not in array_Final:
                array_Final.append(value)
                array_already_processed.append(value)

                # look inside the ".Corr" table
                for value_corr in array_Corr:
                    
                    #building the value to find
                    value_to_find=value["docsymbol"]+"/Corr."
                    if (value_corr["docsymbol"].find(value_to_find)!= -1):
                        array_Final.append(value_corr)
                            
                # look inside the ".Add" table
                for value_add in array_Add:
                    
                    #building the value to find
                    value_to_find=value["docsymbol"]+"/Add."
                    if (value_add["docsymbol"].find(value_to_find)!= -1):
                        array_Final.append(value_add)
                                

        # return the new array
        return array_Final
        
        
        
                

    for record in myRecords:
        
        try :
            itshead= record['head']            
            # Adding itshead content
            dxa_len_curr_head=dxa_counter_head(itshead)
            #if counter + next head +first subhead+first entry_note .. then break as there is no room on the page..
            #cs=dxa_counter_subhead(record["subheading"][0])
            #fe=dxa_counter_entry_note(record['subheading'][0]["entries"][0])
            if (dxa_counter+dxa_len_curr_head+dxa_counter_subhead(record["subheading"][0])+dxa_counter_entry_note(record['subheading'][0]["entries"][0]))>dxa_page_length:#Inches(11-0.61-1-1-1)*1440
                insert_its_head_break(record)
                dxa_counter=0
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(itshead)
            else:
                insert_its_head(record)
                dxa_counter+=dxa_counter_head(itshead)
           
            subheading=record['subheading'] 

            for mysubhead in subheading:

                entries_counter=0
                itssubhead=mysubhead["subhead"]

                # Adding itssubhead content
                
                itsentries1=mysubhead["entries"]
                itsentries=sort_entry(itsentries1)
                dxa_len_entry_note=dxa_counter_entry_note(itsentries[0])
                if (dxa_counter+dxa_counter_subhead(itssubhead)+dxa_counter_entry_note(itsentries[0]))>dxa_page_length:# and entries_counter<=len(itsentries):#Inches(11-0.61-1-1-1)*1440
                    #print(f"dxa_counter BEFORE column break is {dxa_counter}")
                    insert_its_head_break(record)
                    dxa_counter=0
                    insert_its_head_continued(record)
                    dxa_counter+=dxa_counter_head_continued(itshead)
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itshead)                   
                else:
                    insert_its_subhead(itssubhead)
                    dxa_counter+=dxa_counter_subhead(itssubhead)

                formatEntry(itsentries)
                for entry in itsentries:
                    # insertion of the new function
                    
                    entries_counter+=1
                    #Adding itssubhead content
                    if dxa_counter+dxa_counter_entry_note(entry)>dxa_page_length:# and entry["docsymbol"]!=itsentries[-1]["docsymbol"]:#Inches(11-0.61-1-1-1)*1440
                        insert_its_head_break(record)
                        dxa_counter=0
                        insert_its_head_continued(record)
                        dxa_counter+=dxa_counter_head_continued(itshead)
                        insert_entry_note(entry)
                        dxa_counter+=dxa_counter_entry_note(entry)                 
                    else:
                        insert_entry_note(entry)
                        dxa_counter+=dxa_counter_entry_note(entry)
                    

        except:
            itshead=""

    add_page_number(document.sections[0].footer.paragraphs[0])
    return document    



def generateWordDocITPDSL(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection}).sort("sort",1)

    # Creation of the word document

    document = Document()

    # Implementing the "Two columns display" feature



    # # Marging of the document

    # section.top_margin = Cm(2.54)
    # section.bottom_margin = Cm(2.54)
    # section.left_margin = Cm(2.54)
    # section.right_margin = Cm(2.54)
   
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    #font.italic = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Adding the header to the document
    
    header=document.sections[0].header
   
   ################## SORENTRY ###############################################
    
    stlSorentry = document.styles.add_style('sorentry', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    #stlSorentryFont.size = Pt(10)
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.bold = True
    
    pfSorentry = stlSorentry.paragraph_format

    # Line spacing
    
    pfSorentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
        
    ################## SORNORM ###############################################
    
    stlSornorm = document.styles.add_style('sornorm', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
   
    stlSornormFont=stlSornorm.font
    stlSornormFont.name = 'Arial'
    stlSornormFont.size = Pt(8)
    stlSornormFont.bold = True
    
    pfSornorm = stlSornorm.paragraph_format

    # Line spacing
    
    pfSornorm.line_spacing_rule =  WD_LINE_SPACING.SINGLE


    ################## SORNOTE ###############################################
    
    stlSornote = document.styles.add_style('sornote', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSornoteFont=stlSornote.font
    stlSornoteFont.name = 'Arial'
    stlSornoteFont.size = Pt(8)     
    pfSornote = stlSornote.paragraph_format

    # Line spacing
    
    pfSornote.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    ################## SORNOTE-ITALIC ###############################################
    
    stlSornoteIt = document.styles.add_style('sornote_italic', WD_STYLE_TYPE.CHARACTER)
    
    # Font name
    
    stlSornoteFontIt=stlSornoteIt.font
    stlSornoteFontIt.name = 'Arial'
    stlSornoteFontIt.size = Pt(8)    
    stlSornoteFontIt.italic = True 
    #pfSornoteIt = stlSornote.paragraph_format

    # Line spacing
    
    #pfSornoteIt.line_spacing_rule =  WD_LINE_SPACING.SINGLE
     
################## SORNOTE-nonITALIC ###############################################
    
    stlSornoteNonIt = document.styles.add_style('sornote_non_italic', WD_STYLE_TYPE.CHARACTER)
    
    # Font name
    
    stlSornoteFontNonIt=stlSornoteNonIt.font
    stlSornoteFontNonIt.name = 'Arial'
    stlSornoteFontNonIt.size = Pt(8)    
    stlSornoteFontNonIt.italic = False 
    #pfSornoteIt = stlSornote.paragraph_format

    # Line spacing
    
    #pfSornoteIt.line_spacing_rule =  WD_LINE_SPACING.SINGLE
    ################## WRITING THE DOCUMENT ###############################################
    
    # Adding the Header to the document
    
    myRecords=setOfData
    x=0

    p=header.add_paragraph(myTitle.upper(), style='New Heading')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")
    p.add_run("\n")

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'1')

    if (bodysession[0]=="A"):
        p=document.add_paragraph()
        p.add_run("NOTE: ",style="sornote_italic") 
        #run1.text=
        #run1.
        run2=p.add_run("Languages of corrigenda are indicated only when corrigenda are not issued in all six official languages. Documents issued as ")
        run2.style="sornote_non_italic"
        run3=p.add_run("Supplements")
        run3.style="sornote_italic"
        run4=p.add_run(" to the ")
        run4.style="sornote_non_italic"
        run5=p.add_run("Official Records of the General Assembly,")
        run5.style="sornote_italic"
        run6=p.add_run(" Seventy-second Session are so indicated. Information regarding documents bearing the double symbol A/-  and S/-  can be found in the ")
        run6.style="sornote_non_italic"
        run7=p.add_run("Supplements")
        run7.style="sornote_italic"
        run8=p.add_run(" to the ")
        run8.style="sornote_non_italic"
        run9=p.add_run("Official Records of the Security Council.")
        run9.style="sornote_italic"
        run10=p.add_run(" The information provided below is current as of the date this Index is submitted for publication.")
        run10.style="sornote_non_italic"
    if (bodysession[0]=="E"):
        p=document.add_paragraph()
        p.add_run("NOTE: ",style="sornote_italic") 
        run2=p.add_run("Languages of corrigenda are indicated only when corrigenda are not issued in all six languages. Documents issued as ")
        run2.style="sornote_non_italic"
        run3=p.add_run("Supplements")
        run3.style="sornote_italic"
        run4=p.add_run(" to the ")
        run4.style="sornote_non_italic"
        run5=p.add_run("Official Records of the Economic and Social Council,")
        run5.style="sornote_italic"
        run6=p.add_run(" 2018 are also indicated. The information provided below is current as of the date this Index is submitted for publication.")
        run6.style="sornote_non_italic"
    if (bodysession[0]=="S"): 
        p=document.add_paragraph()
        p.add_run("NOTE: ",style="sornote_italic") 
        runs2=p.add_run("Languages of corrigenda are indicated only when corrigenda are not issued in all six official languages. The information provided below is current as of the date this Index is submitted for publication.")
        runs2.style="sornote_non_italic"

    p.add_run("\n")

    p=document.add_paragraph(myRecords[0]["committee"],style="sorentry")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    for record in myRecords:
        
        if x>0:
            p1=document.add_paragraph(record["committee"],style="sorentry")  
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER    
 
        x=x+1
        mySeries = record["series"]
        
        for serie in mySeries:
        
            p2=document.add_paragraph(serie["series"],style="sornorm")
            
            #p2.add_run("\n")
            
            myDocSymbols=serie["docsymbols"]
            
            p3=document.add_paragraph(style="sornote")
            
            for doc in myDocSymbols:
                
                p3.add_run(doc)
                p3.add_run("\n")
                    
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document

def generateWordDocITPMEET(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):

    # Function to keep the header visible for the table

    def set_repeat_table_header(row):

        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})
    setOfData1=myCollection.find({'bodysession': bodysession,'section': paramSection})



    ### calc whether the table of meetings will fit onthe page; breaks the page if it doesn't fit; returns cursorLIne number
    def lines_committee(data1, cursorLine):
        nbMeetings=0
        CURSOR_LINE_MAX=50

        for num in data1["years"]:
            nbMeetings+=len(num["meetings"]) 

        nbYears=len(data1["years"])    
        nbMeetings=nbMeetings+(nbYears-1)*3

        col1_recs=[]
        curr_year=data1["years"][0]
        for num in data1["years"]:
            if curr_year!=num["year"]:
                col1_recs.append((num["year"],'',''))
                col1_recs.append((num["year"],'Date, ',''))
                col1_recs.append((num["year"],'',''))
                curr_year=num["year"]
            for meeting in num["meetings"]:
                col1_recs.append((num["year"],meeting["meetingnum"],meeting["meetingdate"]))
        nbMeetings=len(col1_recs)

        t=col1_recs.pop(0)
        t=col1_recs.pop(0)
        t=col1_recs.pop(0)

        nbMeetings=len(col1_recs)

            
        if nbMeetings>=9:
            nbrRecPerCol1=round(nbMeetings / 3)
            nbrRecPerCol2=round(nbMeetings / 3)
            nbrRecPerCol3=nbMeetings - (nbrRecPerCol1+nbrRecPerCol1)
            nbMeetings=len(col1_recs)
        else:
            nbrRecPerCol1=nbMeetings
            nbrRecPerCol2=nbMeetings
            nbrRecPerCol3=nbMeetings
        
        for num in data1["years"]:   

            if (nbrRecPerCol1 >= nbrRecPerCol3):
                
                if (cursorLine + nbrRecPerCol1 > CURSOR_LINE_MAX):

                    # break the page
                    run=document.add_paragraph().add_run()
                    myBreak = run.add_break(WD_BREAK.PAGE)

                    # reset the cursorLine
                    cursorLine=0

            else :

                if (cursorLine + nbrRecPerCol3 > CURSOR_LINE_MAX):

                    # break the page
                    run=document.add_paragraph().add_run()
                    myBreak = run.add_break(WD_BREAK.PAGE)

                    # reset the cursorLine
                    cursorLine=0

        return col1_recs, cursorLine, nbMeetings, nbrRecPerCol1, nbrRecPerCol2, nbrRecPerCol3
    
    # Creation of the word document
    document = Document()

   
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Adding the header to the document
    
    header=document.sections[0].header

    ################## itssubhead ###############################################
    
    stlItssubHead = document.styles.add_style('itssubhead', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlItsSubHeadFont=stlItssubHead.font
    stlItsSubHeadFont.name = 'Arial'
    stlItsSubHeadFont.size = Pt(8)

    pfItsSubHead = stlItssubHead.paragraph_format

    # Indentation
    
    pfItsSubHead.left_indent = Inches(0)
    
    # Line spacing
    
    pfItsSubHead.line_spacing_rule =  WD_LINE_SPACING.SINGLE
   
   ################## stylemeeting ###############################################
    
    stlSorentry = document.styles.add_style('stylemeeting', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    
    pfSorentry = stlSorentry.paragraph_format

    ################## styleheaderforA ###############################################
    
    stlstyleheaderforA = document.styles.add_style('styleheaderforA', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlstyleheaderforAFont=stlstyleheaderforA.font
    stlstyleheaderforAFont.name = 'Arial'
    stlstyleheaderforAFont.size = Pt(8)
    stlstyleheaderforAFont.bold = True
    
    pfstlstyleheaderforA = stlstyleheaderforA.paragraph_format
    # pfstlstyleheaderforA.space_before = Pt(10)
    # pfstlstyleheaderforA.keep_together = True
    # pfstlstyleheaderforA.keep_with_next = True

    pfstlstyleheaderforA.line_spacing =  1
    

    ################## styletableheader ###############################################
    
    stlSorentry = document.styles.add_style('styletableheader', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    
    pfSorentry = stlSorentry.paragraph_format
    pfSorentry.space_after = Pt(10)
     
    ################## WRITING THE DOCUMENT ###############################################
    
    if (bodysession[0]=="S"):

        # Adding the Header to the document

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'3')

        datas=setOfData

        p=header.add_paragraph(myTitle.upper(), style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")
        
        p=header.add_paragraph("(Symbol: " + datas[0]["symbol"] +")", style='itssubhead')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")
        p.add_run("\n")

        # creation of the table
        table = document.add_table(rows=1, cols=2)

        table.allow_autofit=True

        #table.rows.height_rule = WD_ROW_HEIGHT_RULE.AUTO

        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # set the header visible
        set_repeat_table_header(table.rows[0])

        # Retrieve the first line
        hdr_cells = table.rows[0].cells

        myRun=hdr_cells[0].paragraphs[0].add_run('Meeting')
        myRun.underline=True

        myRun=hdr_cells[1].paragraphs[0].add_run('Date, '+datas[0]["years"][0]["year"])
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        myRun.underline=True

        myRun=hdr_cells[1].paragraphs[0].add_run('\n')

        for data in datas:
            records=data["years"]
            for record in records:
                meetings=record["meetings"]
                for meeting in meetings:
                    row_cells = table.add_row().cells
                    row_cells[0].paragraphs[0].text=meeting["meetingnum"]
                    row_cells[1].paragraphs[0].text=meeting["meetingdate"]
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
    
        # Definition of the size of the column

        widths = (Inches(2), Inches(1.5))
        for row in table.rows:
            for idx,width in enumerate(widths):
                row.cells[idx].width = width

        # Definition of the font

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(0)
                    # Adding styling
                    for run in paragraph.runs:
                        font = run.font
                        font.name="Arial"
                        font.size= Pt(8)
        
        add_page_number(document.sections[0].footer.paragraphs[0])
        return document
    
    '''
    if (bodysession[0]=="Er23"):
        
        datas=setOfData
        datas1=setOfData1

        p=header.add_paragraph(myTitle.upper(), style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")
        p.add_run("\n")
        
        p=header.add_paragraph("(Symbol: " + datas[0]["symbol"] +")", style='itssubhead')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")

        ###########################################################################################
        # Process
        ###########################################################################################
        
        # 0- Settings

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'3')

        nbRecords=0
        nbYears=0
        nbMeetings=0

        # 1- Count the number of records

        for data in datas:
            nbRecords=nbRecords+1
            for year in data["years"]:
                nbYears=nbYears+1
                for meeting in year["meetings"]:
                    nbMeetings=nbMeetings+1

        # 2- Define the number of records per cols
        nbMeeting=1
        nbrRecPerCol0= round(nbMeetings / 3)
        nbrRecPerCol1= round(nbMeetings / 3)
        nbrRecPerCol2= round(nbMeetings / 3)
        nbrRecPerCol3= nbMeetings - (nbrRecPerCol1+nbrRecPerCol2)


        # 3- Define a closure feature

        # 4- Display the record after check of the date a new date should call the closure function

        startGroup=False

        for data1 in datas1:
            for year1 in data1["years"]:

                if startGroup==True:

                    hdr_cells = table.rows[nbMeeting].cells
                    hdr_cells[1].text=" "
                    nbMeeting=nbMeeting+1

                    hdr_cells = table.rows[nbMeeting].cells
                    hdr_cells[1].text="Date, "+ year1["year"]
                    paragraphs = hdr_cells[1].paragraphs
                    
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.underline=True
                    nbMeeting=nbMeeting+1

                    hdr_cells = table.rows[nbMeeting].cells
                    hdr_cells[1].text=" "
                    nbMeeting=nbMeeting+1

                    table.add_row()
                    table.add_row()
                    table.add_row()

                    nbrRecPerCol1=nbrRecPerCol1+3

                for meeting1 in year1["meetings"]:

                    if (nbMeeting<=nbrRecPerCol1):

                        if nbMeeting==1:

                            p=document.add_paragraph()
                            run = p.add_run()
                            
                            table = document.add_table(rows=nbrRecPerCol1+1,cols=2)
                        
                            # set the header visible
                            set_repeat_table_header(table.rows[0])

                            # Retrieve the first line
                            hdr_cells = table.rows[0].cells
                            myRun=hdr_cells[0].paragraphs[0].add_run('Meeting')
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('Date, '+year1["year"])
                            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('\n')
                            #nbMeeting=nbMeeting+1
                            hdr_cells = table.rows[nbMeeting].cells
                            hdr_cells[0].text=meeting1["meetingnum"]
                            hdr_cells[1].text=meeting1["meetingdate"]
                            
                        else :
                            # yls 
                            table = document.add_table(rows=nbrRecPerCol1+1,cols=2)
                            
                            # fill the cell of the table
                            hdr_cells = table.rows[nbMeeting].cells
                            hdr_cells[0].text=meeting1["meetingnum"]
                            hdr_cells[1].text=meeting1["meetingdate"]


                    # if (nbMeeting>nbrRecPerCol1 and nbMeeting<=2*nbrRecPerCol2):
                    if (nbMeeting>nbrRecPerCol1):


                        if nbMeeting==nbrRecPerCol1+1:


                            # column break
                            p=document.add_paragraph()
                            run = p.add_run()
                            myBreak = run.add_break(WD_BREAK.COLUMN)
                            table = document.add_table(rows=nbrRecPerCol1+1,cols=2)
                        
                            # set the header visible
                            set_repeat_table_header(table.rows[0])

                            # Retrieve the first line
                            hdr_cells = table.rows[0].cells
                            myRun=hdr_cells[0].paragraphs[0].add_run('Meeting')
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('Date, '+year1["year"])
                            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('\n')

                            nbMeeting=1
                            nbrRecPerCol1=nbrRecPerCol0
                            hdr_cells = table.rows[nbMeeting].cells
                            hdr_cells[0].text=meeting1["meetingnum"]
                            hdr_cells[1].text=meeting1["meetingdate"]
                            
                        else :

                            # fill the cell of the table
                            hdr_cells = table.rows[nbMeeting].cells
                            hdr_cells[0].text=meeting1["meetingnum"]
                            hdr_cells[1].text=meeting1["meetingdate"]

                    
                    nbMeeting=nbMeeting+1
                
                startGroup=True


        # Apply the font
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    cell.width = Inches(0.7)
                    for paragraph in paragraphs:
                        # Adding styling
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            font = run.font
                            font.name="Arial"
                            font.size= Pt(8)

        add_page_number(document.sections[0].footer.paragraphs[0])
        return document
'''
    if (bodysession[0]=="A") or (bodysession[0]=="E"):   
        datas=setOfData
        datas1=setOfData1

        p=header.add_paragraph(myTitle.upper(), style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")
        

        cursorLine=1
        for data1 in datas:

            
            # 0- Reset the counters 
            
            index0=0
            index1=0
            
            
            line=0        
            overFlowTitle=5
            overFlowContent=3
            createSection=False

            
            # # check the number of lines available
            col1_recs,cursorLine, nbMeetings,nbrRecPerCol1, nbrRecPerCol2, nbrRecPerCol3=lines_committee(data1, cursorLine+overFlowTitle)


            section = document.add_section(WD_SECTION.CONTINUOUS)
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'),'1')
            document.add_paragraph('')

            # writing the title for each record
            #section = document.add_section(WD_SECTION.CONTINUOUS)
            #cursorLine+=1
            p=document.add_paragraph(data1["committee1"],style="styleheaderforA")
            cursorLine+=1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run=p.add_run("\n")
            cursorLine+=1
            if data1["committee2"]:
                p.add_run(data1["committee2"])
                p.add_run("\n")
                cursorLine+=2
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #p.paragraph_format.
        
            if data1["committee1"]!="Credentials Committee":
                z=p.add_run("(Symbol: " + data1["symbol"]+ ")")
            cursorLine+=1
            z.bold=False

            if nbMeetings>=9:
                
                for years in data1["years"]:
                    years['year']      
                section = document.add_section(WD_SECTION.CONTINUOUS)
                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')[0]
                cols.set(qn('w:num'),'3')

                # 3- Display the values using table
                #line=0
                new_page=0
                index0=0
                for index0,entry in enumerate(col1_recs):

                    if index0==0 or index0==nbrRecPerCol1 or index0==(nbrRecPerCol1+nbrRecPerCol2):
                        p=document.add_paragraph()
                        if index0!=0:
                            run=p.add_run()
                            myBreak = run.add_break(WD_BREAK.COLUMN)
                            #new_page=1
                        table = document.add_table(rows=nbrRecPerCol1+overFlowContent,cols=2)
                        table.alignment=WD_TABLE_ALIGNMENT.CENTER
                        line=0
                        # display header
                        hdr_cells = table.rows[0].cells
                        myRun=hdr_cells[0].paragraphs[0].add_run('Meeting')
                        myRun.underline=True
                        myRun=hdr_cells[1].paragraphs[0].add_run('Date, '+ col1_recs[index0][0])
                        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        myRun.underline=True
                        myRun=hdr_cells[1].paragraphs[0].add_run('\n')

                        line+=2
                    if (index0==(nbrRecPerCol1-1) or index0==nbrRecPerCol1) and col1_recs[index0][2]=='' and col1_recs[index0+1][1]=='Date, ':  
                        new_page=1
                        continue
                    elif (index0==(nbrRecPerCol1+nbrRecPerCol2-1) or index0==nbrRecPerCol1+nbrRecPerCol2) and col1_recs[index0][2]=='':  
                        new_page=1
                        continue
                    elif new_page==1 and (index0==(nbrRecPerCol1) or index0==(nbrRecPerCol1+1)):
                        continue
                    elif index0==(nbrRecPerCol1+nbrRecPerCol2-1) and col1_recs[index0][2]=='' and (col1_recs[index0+1][1]=='Date, ' or col1_recs[index0+1][1]==''): 
                        new_page=1
                        continue
                    elif new_page==1 and (index0==nbrRecPerCol1+nbrRecPerCol2 or index0==(nbrRecPerCol1+nbrRecPerCol2+1)):
                        if col1_recs[index0][2]=='':
                            continue
                        else:
                            hdr_cells = table.rows[line].cells
                            hdr_cells[0].text=col1_recs[index0][1]
                            hdr_cells[1].text=col1_recs[index0][2]
                            myActualYear=col1_recs[index0][0]
                            new_page=0
                    elif col1_recs[index0][1]=='Date, ':
                        if not new_page==1:
                            hdr_cells = table.rows[line].cells
                            myRun0=hdr_cells[0].paragraphs[0].add_run('')
                            myRun1=hdr_cells[1].paragraphs[0].add_run(col1_recs[index0][1]+col1_recs[index0][0])
                            myRun1.underline=True
                            myActualYear=col1_recs[index0][0]
                            new_page=0  
                    else:
                        hdr_cells = table.rows[line].cells
                        hdr_cells[0].text=col1_recs[index0][1]
                        hdr_cells[1].text=col1_recs[index0][2]
                        myActualYear=col1_recs[index0][0]
                        new_page=0

                    index0+=1
                    line+=1

            for years in data1["years"]:

                # 2- Define the number of records per cols
                if nbMeetings>=9:
                    pass                                
                                

                else :

                    # creation of one table with 1 column
                    table = document.add_table(rows=nbMeetings+1,cols=2)

                    # 3- Display the values using table
                    for meeting in years["meetings"]:

                        if index1==0:
                            
                            # p=document.add_paragraph()
                            
                            # table creation
                            table.alignment=WD_TABLE_ALIGNMENT.CENTER

                            # display header
                            hdr_cells = table.rows[0].cells
                            cursorLine+=1
                            myRun=hdr_cells[0].paragraphs[0].add_run('Meeting')
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('Date, '+years["year"])
                            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            myRun.underline=True
                            myRun=hdr_cells[1].paragraphs[0].add_run('\n')
                            index1+=1
                            cursorLine+=1

                        # fill the cell of the table
                        hdr_cells = table.rows[index1].cells
                        hdr_cells[0].text=meeting["meetingnum"]
                        hdr_cells[1].text=meeting["meetingdate"]
                        index1+=1
                        cursorLine+=1
                    
                    # add some space after
                    document.add_paragraph('')
                    

        # Apply the font
        widths = (Inches(1.15), Inches(0.85
        ))
        for table in document.tables:
            for row in table.rows:
                for idx,width in enumerate(widths):
                    row.cells[idx].width = width
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        # Adding styling
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = Pt(0)
                        paragraph_format.line_spacing =  1
                        for run in paragraph.runs:
                            font = run.font
                            font.name="Arial"
                            font.size= Pt(8)
        add_page_number(document.sections[0].footer.paragraphs[0])
        return document


def generateWordDocITPAGE(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})

    # Creation of the word document

    document = Document()
   
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    # new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    # Adding the header to the document
    
    header=document.sections[0].header
    

    ################## SORTITLE ###############################################
    
    stlSorentry = document.styles.add_style('sortitle', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.bold = True
    
    pfSorentry = stlSorentry.paragraph_format
    pfSorentry.left_indent = Inches(0)

    # Line spacing
    
    pfSorentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    ################## SORNOTETITLE ###############################################
    
    sornotetitle = document.styles.add_style('sornotetitle', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=sornotetitle.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)

    
    '''
    ################## see ###############################################
    
    see = document.styles.add_style('see', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=see.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.italic=True
    '''
    ################## data ###############################################
    
    data = document.styles.add_style('data', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=data.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.italic=False

    ################## NOTE ###############################################
    
    stlSorentry = document.styles.add_style('note', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    stlSorentryFont.italic = True
    
    pfSorentry = stlSorentry.paragraph_format
    pfSorentry.left_indent = Inches(0)

    # Line spacing
    
    pfSorentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE
   
   ################## SORENTRY ###############################################
    
    stlSorentry = document.styles.add_style('sorentry', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSorentryFont=stlSorentry.font
    stlSorentryFont.name = 'Arial'
    stlSorentryFont.size = Pt(8)
    
    pfSorentry = stlSorentry.paragraph_format

    # Line spacing
    
    pfSorentry.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    ################## see2 ###############################################
    
    stlSee = document.styles.add_style('see2', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSeeFont=stlSee.font
    stlSeeFont.italic=True
    stlSeeFont.name = 'Arial'
    stlSeeFont.size = Pt(8)

    pformat=stlSee.paragraph_format
    pformat.space_after=Pt(5)
    pformat.left_indent=Cm(1.52)
    pformat.first_line_indent=Inches(-0.58)
 
    ################## title ###############################################

    stlTitle = document.styles.add_style('title2', WD_STYLE_TYPE.PARAGRAPH)

    # Font name and size

    stlTitleFont=stlTitle.font
    stlTitleFont.name = 'Arial'
    stlTitleFont.size = Pt(8)

    # others formatting

    pformat=stlTitle.paragraph_format
    pformat.space_before=Pt(8)
    pformat.space_after=Pt(0)
    pformat.left_indent=Inches(0.3)
    pformat.first_line_indent=Inches(-0.3)
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE



    ################## see ###############################################

    see = document.styles.add_style('see', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    seeFont=see.font
    seeFont.name = 'Arial'
    seeFont.size = Pt(8)
    seeFont.italic=True

    pformat=see.paragraph_format
    pformat.space_after=Pt(3)
    pformat.left_indent=Inches(0.6)
    pformat.first_line_indent=Inches(-0.3)
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE


    ################## see2a ###############################################

    see2a = document.styles.add_style('see2a', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see2aFont=see2a.font
    see2aFont.name = 'Arial'
    see2aFont.size = Pt(8)
    see2aFont.italic=True

    pformat2=see2a.paragraph_format
    pformat2.space_after=Pt(3)
    pformat2.left_indent=Inches(0.6)
    pformat2.first_line_indent=Inches(-0.3)
    pformat2.line_spacing_rule = WD_LINE_SPACING.SINGLE


    ################## see2b ###############################################

    see2b = document.styles.add_style('see2b', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see2bFont=see2b.font
    see2bFont.name = 'Arial'
    see2bFont.size = Pt(8)
    see2bFont.italic=True

    pformat2=see2b.paragraph_format
    pformat2.space_after=Pt(3)
    pformat2.left_indent=Inches(0.7)
    pformat2.first_line_indent=Inches(-0.1)
    pformat2.line_spacing_rule = WD_LINE_SPACING.SINGLE



    ################## seealso ###############################################

    seealso = document.styles.add_style('seealso', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    seealsoFont=seealso.font
    seealsoFont.name = 'Arial'
    seealsoFont.size = Pt(8)
    seealsoFont.italic=True

    pformatseealso=seealso.paragraph_format
    pformatseealso.space_after=Pt(3)
    pformatseealso.left_indent=Inches(0.8)
    pformatseealso.first_line_indent=Inches(-0.5)
    pformatseealso.line_spacing_rule = WD_LINE_SPACING.SINGLE


    ################## see_also_subjs ###############################################

    see_also_subjs = document.styles.add_style('see_also_subjs', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see_also_subjsFont=see_also_subjs.font
    see_also_subjsFont.name = 'Arial'
    see_also_subjsFont.size = Pt(8)
    see_also_subjsFont.italic=True

    pformatsee_also_subjs=see_also_subjs.paragraph_format
    pformatsee_also_subjs.space_after=Pt(3)
    pformatsee_also_subjs.left_indent=Inches(1.1)
    #pformatsee_also_subjs.first_line_indent=Inches(-0.5)
    pformatsee_also_subjs.line_spacing_rule = WD_LINE_SPACING.SINGLE

    ################## see3 ###############################################

    see3 = document.styles.add_style('see3', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see3Font=see3.font
    see3Font.name = 'Arial'
    see3Font.size = Pt(8)
    see3Font.italic=True

    pformat3=see3.paragraph_format
    pformat3.space_after=Pt(5)
    pformat3.left_indent=Inches(0.9)
    pformat3.first_line_indent=Inches(-0.3)
    pformat3.line_spacing_rule = WD_LINE_SPACING.SINGLE
    ################## see4 ###############################################

    see4 = document.styles.add_style('see4', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see4Font=see4.font
    see4Font.name = 'Arial'
    see4Font.size = Pt(8)
    see4Font.italic=True

    pformat4=see4.paragraph_format
    pformat4.space_after=Pt(5)
    pformat4.left_indent=Inches(1.2)
    pformat4.first_line_indent=Inches(-0.3)
    pformat4.line_spacing_rule = WD_LINE_SPACING.SINGLE

    ################## see5 ###############################################

    see5 = document.styles.add_style('see5', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see5Font=see5.font
    see5Font.name = 'Arial'
    see5Font.size = Pt(8)
    see5Font.italic=True

    pformat5=see5.paragraph_format
    pformat5.space_after=Pt(5)
    pformat5.left_indent=Inches(0.8)
    #pformat4.first_line_indent=Inches(-0.3)
    pformat5.line_spacing_rule = WD_LINE_SPACING.SINGLE

################## see5E ###############################################

    see5E = document.styles.add_style('see5E', WD_STYLE_TYPE.PARAGRAPH)

    # Font name

    see5EFont=see5E.font
    see5EFont.name = 'Arial'
    see5EFont.size = Pt(8)
    see5EFont.italic=True

    pformat5E=see5E.paragraph_format
    pformat5E.space_after=Pt(5)
    pformat5E.left_indent=Inches(1.1)
    pformat5E.first_line_indent=Inches(-0.3)
    pformat5E.line_spacing_rule = WD_LINE_SPACING.SINGLE

    ################## subagenda2 ###############################################
    #style for first subagenda (a), (b)
    stlSubagenda2 = document.styles.add_style('subagenda2', WD_STYLE_TYPE.PARAGRAPH)

    # Font name and size

    stlSubagenda2Font=stlSubagenda2.font
    stlSubagenda2Font.name = 'Arial'
    stlSubagenda2Font.size = Pt(8)

    # others formatting

    pformat=stlSubagenda2.paragraph_format
    pformat.space_before=Pt(0)
    pformat.space_after=Pt(2)
    pformat.left_indent=Inches(0.6)
    pformat.first_line_indent=Inches(-0.3)
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE

    ################## subagenda3 ###############################################
    #style for aa is (a), bb is (b), etc 
    stlSubagenda3 = document.styles.add_style('subagenda3', WD_STYLE_TYPE.PARAGRAPH)

    # Font name and size

    stlSubagenda3Font=stlSubagenda3.font
    stlSubagenda3Font.name = 'Arial'
    stlSubagenda3Font.size = Pt(8)

    # others formatting

    pformat=stlSubagenda3.paragraph_format
    pformat.space_before=Pt(0)
    pformat.space_after=Pt(2)
    pformat.left_indent=Inches(0.9)
    pformat.first_line_indent=Inches(-0.3)
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE

    ################## seesubagenda3 ###############################################

    stlseeSubagenda3 = document.styles.add_style('seesubagenda3', WD_STYLE_TYPE.PARAGRAPH)

    # Font name and size

    stlSubagenda3Font=stlseeSubagenda3.font
    stlSubagenda3Font.name = 'Arial'
    stlSubagenda3Font.size = Pt(8)

    # others formatting

    pformat=stlseeSubagenda3.paragraph_format
    pformat.space_before=Pt(0)
    pformat.space_after=Pt(2)
    pformat.left_indent=Inches(1.2)
    pformat.first_line_indent=Inches(-0.3)
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE


    
     
    ################## WRITING THE DOCUMENT ###############################################
    
    # Adding the Header to the document

    ########### GENERAL ASSEMBLY ##########

    def key_title(myAgendaText):
            if myAgendaText["title"]=="":
                #return (2,myAgendaText["title"])
                return 2
            else:
                #return (1,myAgendaText["title"])
                return 1

    if (bodysession[0]=="A"):

        ########### GENERAL ASSEMBLY ##########

        myRecords=setOfData
        preventDoubleValue=""


        for record in myRecords:  

            # Display management
            
            if record["heading"]=="AGENDA":

                p=header.add_paragraph(myTitle.upper(), style='New Heading')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                section = document.sections[0]
                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')[0]
                cols.set(qn('w:num'),'2')

                # Selection of the agenda object
                for myAgenda in record["agendas"]:

                    # Selection of the agendanum value
                    myAgendaNum=myAgenda["agendanum"]
                    prev_my_subagenda=""
                    # Selection of the text object
                    alreadyDisplayText=False
                    alreadyDisplayNum=False
                    alreadyDisplaySee=False
                    alreadyDisplaySeeAlso=False
                    alreadyDisplaySeeAlso1=False
                    saveTitle=""
                    seeAlso1=False
                    for myText in myAgenda["text"]:
                        
                        # Selection of the subagenda value
                        if myText["subagenda"]:
                            mySubAgenda=myText["subagenda"].strip()
                            alreadyDisplaySee=False
                            alreadyDisplaySeeAlso=False
                        else:
                            mySubAgenda=""

                        # Selection of the agendatext value
                        mySortedText=sorted(myText["agendatext"], key=key_title)

                        for i,myAgendaText in enumerate(mySortedText):

                            # Selection of the title value

                            if myAgendaText["title"]!="":
                                myTitle=myAgendaText["title"].strip()
                            else:
                                myTitle=""
                                #alreadyDisplaySee=True

                            if alreadyDisplayNum==False and myTitle!="": #main agenda num e.g. 1., 2., 3., ...
                                p=document.add_paragraph('{}.{}'.format(myAgendaNum,' '*(8-2*len(myAgendaNum)+1)),style="title2")
                                if myAgendaNum=="1":
                                    p.paragraph_format.space_before=Pt(0)
                                run=p.add_run(myTitle)
                                alreadyDisplayNum=True
                                
                            elif len(myAgenda["text"])!=1: #subagenda 
                                if saveTitle!=myTitle and myTitle!="":
                                    if prev_my_subagenda!=mySubAgenda:#subagenda (a)
                                        p=document.add_paragraph("("+  mySubAgenda    +")     ",style="subagenda2")   
                                        run=p.add_run(myTitle)
                                        saveTitle=myTitle
                                        seeAlso1=True
                                        
                                    else:# double subagenda as in agenda_num 98 in A/74 (aa), (bb)
                                        p=document.add_paragraph("("+  mySubAgenda    +")     ",style="subagenda3")   
                                        run=p.add_run(myTitle)
                                        saveTitle=myTitle
                                        seeAlso1=True
                            
                            # Selection of the subject value
                            alreadyDisplayMultipleSee=False
                            alreadyDisplayMultipleb=False
                            index=0
                        
                            for mySubject in myAgendaText["subjects"]:
                                if mySubject:
                                    #my main agenda See when no subagendas exist: 
                                    if alreadyDisplaySee==False and myAgendaText['title']!="":
                                        if len(myAgenda["text"])==1:
                                            p=document.add_paragraph(""+"See:  ",style='see2a') 
                                            run=p.add_run(mySubject.strip())
                                            run.font.italic=False
                                            alreadyDisplaySee=True

                                        else:
                                            # See for subagenda 
                                            if prev_my_subagenda!=mySubAgenda and mySubAgenda!=""and myAgendaText['title']!="":
                                                p=document.add_paragraph("See:  ",style="see3")
                                                run=p.add_run(mySubject.strip()) 
                                                run.font.italic=False
                                                alreadyDisplaySee=True
                                                prev_my_subagenda=mySubAgenda
                                            elif mySubAgenda!=""and myAgendaText['title']!="":# See for the (aa), (bb), (cc),...
                                                    p=document.add_paragraph("See:  ",style="see4")
                                                    run=p.add_run(mySubject.strip()) 
                                                    run.font.italic=False
                                                    alreadyDisplaySee=True
                                            
                                            # See under the main title when there are subagendas as well - subjects
                                            if mySubAgenda=="":
                                                p=document.add_paragraph("See:  ",style="see2a")
                                                run=p.add_run(mySubject.strip()) 
                                                run.font.italic=False
                                                alreadyDisplaySee=True                                           
                                    else:
                                        
                                        # normal scenario
                                        if len(myAgenda["text"])==1 and myText["agendatext"][0]["subjects"][0]=="":
                                            ####### yls
                                            p=document.add_paragraph(""+mySubject.strip(),style="see3")  
                                            alreadyDisplayMultipleb=True
                                            alreadyDisplaySee=True
                                        #See also        
                                        if len(mySortedText)==2 and myAgendaText['title']=="" and myAgendaText['subjects']!="":
                                            if mySubAgenda!="": #subagenda see also
                                                see_also_style="see3"#See also: 1st subject
                                                see_also_subj="see_also_subjs"#subjects
                                            else: # main agenda see also
                                                see_also_style="seealso"
                                                see_also_subj="see5"
                                            if alreadyDisplaySeeAlso==False:
                                                p=document.add_paragraph("See also: ",style=see_also_style)
                                                run=p.add_run(mySubject.strip()) 
                                                run.font.italic=False   
                                                alreadyDisplaySeeAlso=True
                                            else:
                                                p=document.add_paragraph("",style=see_also_subj)
                                                run=p.add_run(mySubject.strip())
                                                run.font.italic=False 

                                            continue

                                else : # no subject
                                    if len(myAgenda["text"])==1:
                                        #run=p.add_run("\n")
                                        p.paragraph_format.space_after = Pt(5)
                                    #pass 


            if record["heading"]=="OTHER MATTERS INCLUDED IN THE INDEX":   

                myBreak = p.add_run().add_break(WD_BREAK.PAGE)

                paragraph = header.paragraphs[0]

                # Choosing the top most sectin of the page 
                section = document.sections[0] 
                
                # Selecting the header 
                header = section.header 
                
                # Selecting the paragraph already present in 
                # the header section 
                header_para = header.paragraphs[0] 

                # Selection of the agenda object
                for myAgenda in record["agendas"]:

                    for myText in myAgenda["text"]:

                        for myAgendaText in myText["agendatext"]:

                            for mySubject in myAgendaText["subjects"]:

                                p=document.add_paragraph(mySubject,style="title2")    



        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # add_page_number(document.sections[0].footer.paragraphs[0])
        # return document

    # add_page_number(document.sections[0].footer.paragraphs[0])
    # return document


    ### SECURITY COUNCIL ##########

    if (bodysession[0]=="S"):

        myRecords=setOfData

        p=document.add_paragraph(myTitle.upper(), style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'1')
        
        p1=document.add_paragraph("", style='sornotetitle')
        p1.add_run("The Council's practice is to adopt at each meeting, on the basis of a provisional agenda circulated in advance, the agenda for that meeting. At subsequent meetings an item may appear in its original form or with the addition of such sub-items as the Council may decide to include. Once included in the agenda, an item remains on the list of matters of which the Council is seized, until the Council agrees to its removal.")
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW

        p2=document.add_paragraph("", style='sornotetitle')
        p2.add_run('The agenda as adopted for each meeting in XXXX will be found in the ')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW

        p2.add_run('Official Records of the Security Council, Seventy-third Year').italic=True
        
        p2.add_run(' (S/PV.8152-8439). A list of weekly summary statements of matters of which the Security Council is seized, and on the stage reached in their consideration, submitted by the Secretary-General under rule 11 of the provisional rules of procedure of the Security Council, appears in the Subject index under the heading "UN. SECURITY COUNCIL (XXXX)–AGENDA".')

        p3=document.add_paragraph("", style='sornotetitle')
        p3.add_run('Listed below are the matters considered by, or brought to the attention of the Council during XXXX. They are arranged alphabetically by the subject headings under which related documents are to be found in the Subject index.')
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
        p3.add_run("\n")
    
        p.add_run("\n")

        p=document.add_paragraph(""+myRecords[0]["heading"],style="sortitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p=document.add_paragraph('NOTE: Subject headings under which documentation related to agenda items is listed \n    in the Subject index appear in capital letters following the title of the item.',style='note')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run("\n")
        p.add_run("\n")

        p=document.add_paragraph('',style='sorentry')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for record in myRecords[0]["agendas"]:

            p.add_run("[To be completed manually]")
            p.add_run("\n")
            run=p.add_run("See    ")
            run.font.italic=True
            run=p.add_run(record["subject"])
            run.font.italic=False
            p.add_run("\n")
            p.add_run("\n")

            
        ### SECOND RECORD ##########
        p=document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # run = p.add_run()
        # run.add_break(WD_BREAK.PAGE)

        p=document.add_paragraph(myRecords[1]["heading"],style="sortitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p=document.add_paragraph('NOTE: These items were not discussed by the Council',style='note')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run("\n")
        p.add_run("\n")

        p=document.add_paragraph('',style='sorentry')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for record in myRecords[1]["agendas"]:

            p.add_run("[To be completed manually]")
            p.add_run("\n")
            run=p.add_run("See    ")
            run.font.italic=True
            run=p.add_run(record["subject"])
            run.font.italic=False
            p.add_run("\n")
            p.add_run("\n")

        # add_page_number(document.sections[0].footer.paragraphs[0])
        # return document

    ########## ECOSOC ##########

    if (bodysession[0]=="E"):
        print(f"body is {bodysession[0]}")
        myRecords=setOfData
        preventDoubleValue=""
        for record in myRecords:  
            # Display management
            if record["heading"]=="AGENDA":
                p=header.add_paragraph(myTitle.upper(), style='New Heading')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                section = document.sections[0]
                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')[0]
                cols.set(qn('w:num'),'1')

                # Selection of the agenda object
                for myAgenda in record["agendas"]:

                    # Selection of the agendanum value
                    myAgendaNum=myAgenda["agendanum"]
                    prev_my_subagenda=""
                    # Selection of the text object
                    alreadyDisplayText=False
                    alreadyDisplayNum=False
                    alreadyDisplaySee=False
                    alreadyDisplaySeeAlso=False
                    alreadyDisplaySeeAlso1=False
                    saveTitle=""
                    seeAlso1=False
                    for myText in myAgenda["text"]:
                        
                        # Selection of the subagenda value
                        if myText["subagenda"]:
                            mySubAgenda=myText["subagenda"].strip()
                            alreadyDisplaySee=False
                            alreadyDisplaySeeAlso=False
                        else:
                            mySubAgenda=""

                        # Selection of the agendatext value
                        mySortedText=sorted(myText["agendatext"], key=key_title)

                        for i,myAgendaText in enumerate(mySortedText):

                            # Selection of the title value

                            if myAgendaText["title"]!="":
                                myTitle=myAgendaText["title"].strip()
                            else:
                                myTitle=""
                                
                                #alreadyDisplaySee=True

                            if alreadyDisplayNum==False and myTitle!="": #main agenda num e.g. 1., 2., 3., ...
                                p=document.add_paragraph('{}.{}'.format(myAgendaNum,' '*(8-2*len(myAgendaNum)+1)),style="title2")
                                if myAgendaNum=="1":
                                    p.paragraph_format.space_before=Pt(0)
                                run=p.add_run(myTitle)
                                alreadyDisplayNum=True
                                
                            elif len(myAgenda["text"])!=1: #subagenda 
                                if saveTitle!=myTitle and myTitle!="":
                                    if prev_my_subagenda!=mySubAgenda:#subagenda (a)
                                        p=document.add_paragraph("("+  mySubAgenda    +")     ",style="subagenda2")   
                                        run=p.add_run(myTitle)
                                        saveTitle=myTitle
                                        seeAlso1=True
                                        
                                    else:# double subagenda as in agenda_num 98 in A/74 (aa), (bb)
                                        p=document.add_paragraph("("+  mySubAgenda    +")     ",style="subagenda3")   
                                        run=p.add_run(myTitle)
                                        saveTitle=myTitle
                                        seeAlso1=True
                            
                            # Selection of the subject value
                            alreadyDisplayMultipleSee=False
                            alreadyDisplayMultipleb=False
                            index=0
                        
                            for mySubject in myAgendaText["subjects"]:
                                if mySubject:
                                    #my main agenda See when no subagendas exist: 
                                    if alreadyDisplaySee==False and myAgendaText['title']!="":
                                        if len(myAgenda["text"])==1:
                                            p=document.add_paragraph(""+"See:  ",style='see2a') 
                                            run=p.add_run(mySubject.strip())
                                            run.font.italic=False
                                            alreadyDisplaySee=True

                                        else:
                                            # See for subagenda 
                                            if prev_my_subagenda!=mySubAgenda and mySubAgenda!=""and myAgendaText['title']!="":
                                                p=document.add_paragraph("See:  ",style="see3")
                                                run=p.add_run(mySubject.strip()) 
                                                run.font.italic=False
                                                alreadyDisplaySee=True
                                                prev_my_subagenda=mySubAgenda
                                            elif mySubAgenda!=""and myAgendaText['title']!="":# See for the (aa), (bb), (cc),...
                                                    p=document.add_paragraph("See:  ",style="see4")
                                                    run=p.add_run(mySubject.strip()) 
                                                    run.font.italic=False
                                                    alreadyDisplaySee=True
                                            
                                            # See under the main title when there are subagendas as well - subjects
                                            if mySubAgenda=="":
                                                p=document.add_paragraph("See:  ",style="see2a")
                                                run=p.add_run(mySubject.strip()) 
                                                run.font.italic=False
                                                alreadyDisplaySee=True                                           
                                    else:
                                        
                                        # normal scenario
                                        #if len(myAgenda["text"])==1 and myText["agendatext"][0]["subjects"][0]=="":
                                            ####### yls
                                        #   p=document.add_paragraph(""+mySubject.strip(),style="see3")  
                                        #    alreadyDisplayMultipleb=True
                                        #    alreadyDisplaySee=True
                                        #See also        
                                        if len(mySortedText)==2 and myAgendaText['title']=="" and myAgendaText['subjects']!="":
                                            if mySortedText[0]["subjects"]==['']:
                                                if mySubAgenda!="": #subagenda see also
                                                    see_also_style="see3"#See also: 1st subject
                                                    see_also_subj="see4"#subjects
                                                else: # main agenda see also
                                                    see_also_style="see2a"
                                                    see_also_subj="see2b"
                                                if alreadyDisplaySeeAlso==False:
                                                    p=document.add_paragraph("See:  ",style=see_also_style)
                                                    run=p.add_run(mySubject.strip()) 
                                                    run.font.italic=False   
                                                    alreadyDisplaySeeAlso=True
                                                else:
                                                    p=document.add_paragraph("",style=see_also_subj)
                                                    run=p.add_run(mySubject.strip())
                                                    run.font.italic=False 
                                            else:
                                                if mySubAgenda!="": #subagenda see also
                                                    see_also_style="see3"#See also: 1st subject
                                                    see_also_subj="see_also_subjs"#subjects
                                                else: # main agenda see also
                                                    see_also_style="seealso"
                                                    see_also_subj="see5"
                                                if alreadyDisplaySeeAlso==False:
                                                    p=document.add_paragraph("See also: ",style=see_also_style)
                                                    run=p.add_run(mySubject.strip()) 
                                                    run.font.italic=False   
                                                    alreadyDisplaySeeAlso=True
                                                else:
                                                    p=document.add_paragraph("",style=see_also_subj)
                                                    run=p.add_run(mySubject.strip())
                                                    run.font.italic=False 
                                                continue


                                else : # no subject
                                    if len(myAgenda["text"])==1:
                                        #run=p.add_run("\n")
                                        p.paragraph_format.space_after = Pt(5)
                                    #pass 

            
            if record["heading"]=="OTHER MATTERS INCLUDED IN THE INDEX":   

                myBreak = p.add_run().add_break(WD_BREAK.PAGE)

                paragraph = header.paragraphs[0]

                # Choosing the top most sectin of the page 
                section = document.sections[0] 
                
                # Selecting the header 
                header = section.header 
                
                # Selecting the paragraph already present in 
                # the header section 
                header_para = header.paragraphs[0] 

                # Selection of the agenda object
                for myAgenda in record["agendas"]:

                    for myText in myAgenda["text"]:

                        for myAgendaText in myText["agendatext"]:

                            for mySubject in myAgendaText["subjects"]:

                                p=document.add_paragraph(mySubject,style="title2")    




        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    add_page_number(document.sections[0].footer.paragraphs[0])
    #document.save('bgitpage'+str(math.floor(datetime.datetime.utcnow().timestamp()))+'.doc')
    return document


def generateWordDocITPVOT(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    config_coll=myDatabase['itp_config']
    myTitle=paramTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})
    setOfData1=myCollection.find({'bodysession': bodysession,'section': paramSection})
    setOfData2=myCollection.find({'bodysession': bodysession,'section': paramSection})
    setOfData3=myCollection.find({'bodysession': bodysession,'section': paramSection})
    lstVotes=list(setOfData3)
    # Creation of the word document

    document = Document()



    def set_cell_border(cell, **kwargs):
        '''
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        '''
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def set_cell_margins(cell, **kwargs):
        '''
        cell:  actual cell instance you want to modify

        usage:

            set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

        provided values are in twentieths of a point (1/1440 of an inch).
        read more here: http://officeopenxml.com/WPtableCellMargins.php
        '''
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in [
            "top",
            "start",
            "bottom",
            "end",
        ]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)


    def set_repeat_table_header(row):

        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    # Adding the header to the document
    
    header=document.sections[0].header
    header.orientation=WD_ORIENT.LANDSCAPE

    section = document.sections[0]
    section.orientation=WD_ORIENT.LANDSCAPE
    new_height = section.page_width
    section.page_width = section.page_height
    section.page_height = new_height
   
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    pfNew_heading_style = new_heading_style.paragraph_format
    pfNew_heading_style.alignment=WD_ALIGN_PARAGRAPH.CENTER
    

    ################## SUBTITLE ###############################################
    
    stlSubtitle = document.styles.add_style('subtitle', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlSubtitleFont=stlSubtitle.font
    stlSubtitleFont.name = 'Arial'
    stlSubtitleFont.size = Pt(8)
    
    pfSubtitle = stlSubtitle.paragraph_format
    pfSubtitle.left_indent = Inches(0.80)
    pfSubtitle.alignment=WD_ALIGN_PARAGRAPH.CENTER

    # Line spacing
    
    pfSubtitle.line_spacing_rule =  WD_LINE_SPACING.SINGLE

    ################## CONTENT ###############################################
    
    stlContent = document.styles.add_style('content', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font name
    
    stlContentFont=stlContent.font
    stlContentFont.name = 'Arial'
    stlContentFont.size = Pt(8)
    
    pfContent = stlContent.paragraph_format

    # Line spacing
    
    pfContent.line_spacing_rule =  WD_LINE_SPACING.DOUBLE
     
    ################## WRITING THE DOCUMENT ###############################################
    
    # Adding the Header to the document
    
    myRecords=list(setOfData)
    myRecords1=list(setOfData1)
    myRecords2=list(setOfData2)

    ### scenario for security council ##########

    if (bodysession[0]=="S"):
        
        # Header

        p=header.add_paragraph("\t"+myTitle.upper(), style='New Heading')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=header.add_paragraph("",style='subtitle')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n")
        p.add_run("Votes are as indicated in the provisional verbatim records of the Security Council,")
        p.add_run("\n")
        p.add_run('seventy-third year, 2018. The following symbols are used to indicate how each member voted:')
        p.add_run("\n")
        p.add_run('   Y	      Voted   Yes')
        p.add_run("\n")
        p.add_run('  N	     Voted   No')
        p.add_run("\n")
        p.add_run(' A	    Abstained')
        p.add_run("\n")
        p.add_run("\t"+'     NP    Not Participating')
        p.add_run("\n")
        p.add_run('Resolutions adopted without vote are indicated by a blank space')

        # Retrieving the number of records
        recordNumber = 0

        # Country list
        countries=[]     
        for datax in myRecords2[0]["votelist"]:
            countries.append(datax["memberstate"])

        recordNumber=len(lstVotes)
        #for record in myRecords:
        #    recordNumber+=1
       

        # creation of the table
        table = document.add_table(rows=1,cols=16)
        print("tableau cree 0")
        table.style = 'TableGrid'

        first=True
        myRow = 0
        myCol = 0
        
        # Variable for managing the number of column to build
        columnToBuild=recordNumber
        
        for record1 in myRecords1:     
            
            # Check if we have already 15 records

            if myRow==15:
            
                myRow=0

                if myCol==16:

                    p=document.add_paragraph()

                    myBreak = p.add_run().add_break(WD_BREAK.PAGE)
                    
                    columnToBuild=(columnToBuild-15)
                    
                    if columnToBuild>=15 :
                    
                        table = document.add_table(rows=1,cols=16)

                        
                    if (columnToBuild<15):
                        
                        table = document.add_table(rows=1,cols=columnToBuild+1) 
                        
                        table.alignment = WD_TABLE_ALIGNMENT.LEFT


                    table.style = 'TableGrid'

                    myCol=0
 
            if myRow==0 and myCol==0:

                paragraph=table.cell(myRow,myCol).paragraphs[0]

                paragraph.text="S/RES/-"                
                
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Inches(0.08)
                paragraph_format.space_after = Inches(0.08)    
                
                table.rows[0].cells[0].width = Inches(1.5)

                run = paragraph.runs

                font = run[0].font

                font.size= Pt(8)

                font.name = "Arial"

                # loop for the countries name to display in column

                for country in countries:

                    table.add_row()
                    
                    myRow=myRow+1
                    
                    paragraph=table.cell(myRow,myCol).paragraphs[0]
                    
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Inches(0.08)
                    paragraph_format.space_after = Inches(0.08)

                    paragraph.text=country

                    run = paragraph.runs

                    font = run[0].font

                    font.size= Pt(8)

                    font.name = "Arial"

                    # size of the first column

                    table.cell(myRow,myCol).width=Inches(2.0)

                
                myRow=0
                myCol=myCol + 1

            if myRow==0 and myCol!=0:

                paragraph=table.cell(myRow,myCol).paragraphs[0]

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                
                #add_hyperlink(paragraph,record1["resnum"],"http://undocs.org/s/res/"+record1["docsymbol"])
                add_hyperlink(paragraph,record1["resnum"],"http://undocs.org/"+record1["docsymbol"])
               
                
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Inches(0.08)
                paragraph_format.space_after = Inches(0.08)  
                
                table.rows[myRow].cells[myCol].width = Inches(0.53)   

                run = paragraph.runs

                font = run[0].font

                font.size= Pt(8)

                font.name = "Arial"


                for data in record1["votelist"]:
                    
                    myRow=myRow+1

                    paragraph=table.cell(myRow,myCol).paragraphs[0]

                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    paragraph.text=data["vote"]  
                    
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Inches(0.08)
                    paragraph_format.space_after = Inches(0.08)   
                    
                    table.rows[myRow].cells[myCol].width = Inches(0.53)  

                    run = paragraph.runs

                    font = run[0].font

                    font.size= Pt(8)

                    font.name = "Arial"
                    

            myCol = myCol + 1

        add_page_number(document.sections[0].footer.paragraphs[0])
        return document


    if (bodysession[0]=="A" and not bodysession=="A/10emsp"):
            #start = time.perf_counter()
        # Retrieving the number of records
        recordNumber = 0

        #bg - 
        myRecords=lstVotes
        myRecords1=lstVotes
        myRecords2=lstVotes
        for record in myRecords:
            recordNumber+=1
        

        # two columns display

        # Country list
        countries=[]     
        for datax in myRecords2[0]["votelist"]:
            countries.append(datax["memberstate"])

        # Adding the Header to the document

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
        
      

        # set some variables
        myRow = 0
        myCol = 0
        print(f"bodysession is {bodysession} and paramSection is {paramSection}")
        try:    
            smyNumColum=config_coll.find_one({'bodysession':bodysession,'section':paramSection,'type':'section'})['no_of_columns']
            scell_width=config_coll.find_one({'bodysession':bodysession,'section':paramSection,'type':'section'})['cell_width']
            sleft_margin=config_coll.find_one({'bodysession':bodysession,'section':paramSection,'type':'section'})['left_margin']
            sright_margin=config_coll.find_one({'bodysession':bodysession,'section':paramSection,'type':'section'})['right_margin']
            print(f"number of columns is {smyNumColum}")
        except:
            smyNumColum='10'
            scell_width='0.38'
            sleft_margin='0.5'
            sright_margin='0.5'

          # Marging of the document

        section.top_margin = Inches(0.81)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(float(sleft_margin))
        section.right_margin = Inches(float(sright_margin)) 
        myNumColum=int(smyNumColum)
        myNumTable=0
        #myNumRow=193 # number of countries that have a vote
        myNumRow=len(countries)
        myIndex=0
        

        # generate the number of table
        myNumTable=math.ceil(recordNumber/int(myNumColum))

        # creation of the table
        table = document.add_table(rows=myNumRow+3, cols=int(myNumColum))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_cells=table._cells
        # write the memberstates in the first (0) column
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        myRow+=1
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        myRow+=1
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        
        for country in countries:
            myRow+=1
            table_cells[myRow*myNumColum].paragraphs[0].text=country
            
        myRow=0


        # Flip
        myFlip=0

        # set the header visible
        set_repeat_table_header(table.rows[0])
        set_repeat_table_header(table.rows[1])
        set_repeat_table_header(table.rows[2])

        
        myOneTimeExec=False
        vlCounter=0
        
        ts1=time.time()
        while myIndex<recordNumber:
            # for 1 row which is a header populate resolution numbers. Votes start from idx=3 and and with idx=195
            
            vlLen=len(myRecords1[myIndex]["votelist"])
            myRow=0
            
            
            if myRow == 0:
                if myOneTimeExec==False:
                    myResCounter=0
                    table_cells[myRow+myCol*myNumColum].paragraphs[0].text=myRecords1[myIndex]["docsymbol"][:9]+"-"
                    myOneTimeExec=True
                # adding two additional row
                
                if myFlip==0:
                    myFlip=1
                    #table_cells[myCol+1+myNflask runumColum].paragraphs[0].text=myRecords1[myIndex]["resnum"]
                    paragraph=table_cells[myCol+1+myNumColum].paragraphs[0]
                    add_hyperlink(paragraph,myRecords1[myIndex]["resnum"],"http://undocs.org/"+myRecords1[myIndex]["docsymbol"])
                    myResCounter+=1   
                else:
                    myFlip=0
                    #table_cells[myCol+1].paragraphs[0].text=myRecords1[myIndex]["resnum"]
                    paragraph=table_cells[myCol+1].paragraphs[0]
                    add_hyperlink(paragraph,myRecords1[myIndex]["resnum"],"http://undocs.org/"+myRecords1[myIndex]["docsymbol"])
                    myResCounter+=1
                    

                myRow+=1
                myRow+=1
                myRow+=1
            #write 1 column of votes
            
            for myVoteList in myRecords1[myIndex]["votelist"]: 
                table_cells[(myRow)*(myNumColum)+myCol+1].paragraphs[0].text=myVoteList["vote"]
                myRow+=1
                vlCounter+=1    
                if vlCounter==vlLen and myCol+1==myNumColum-1:
                    # check if we need to create a new table
                        myCol=0
                        myRow=0
                        myIndex+=1
                        vlCounter=0
                        #create a new table
                        p=document.add_paragraph()
                        run = p.add_run()
                        myBreak = run.add_break(WD_BREAK.PAGE)
                        table = document.add_table(rows=myNumRow+3, cols=myNumColum)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table_cells=table._cells
                        # write the memberstates in the first (0) column
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
                        myRow+=1
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
                        myRow+=1
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""


                        # set the header visible
                        set_repeat_table_header(table.rows[0])
                        set_repeat_table_header(table.rows[1])
                        set_repeat_table_header(table.rows[2])

                        
                        for country in countries:
                            myRow+=1
                            table_cells[myRow*myNumColum].paragraphs[0].text=country
                        
                        myOneTimeExec=False

                    #write next column 
                elif vlCounter==vlLen:    
                        myCol+=1
                        myRow=0
                        myIndex+=1
                        vlCounter=0
                        
                    
            
        stl_itp_vot_cell= document.styles.add_style('itpvotcell', WD_STYLE_TYPE.PARAGRAPH)
        stl_itp_vot_cell_font=stl_itp_vot_cell.font
        stl_itp_vot_cell_font.name = 'Arial'
        stl_itp_vot_cell_font.size = Pt(8)
        stl_itp_vot_cell_font.bold = False
        #pf_stl_itp_vot_cell = stl_itp_vot_cell.paragraph_format

        for table in document.tables:
            #table.autofit = False
            #table.allow_autofit = False
            table_cells=table._cells
            for cell in table_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.style=stl_itp_vot_cell
                    ##pass
                    # Adding styling
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.line_spacing =  1
                    paragraph_format.left_indent = Inches(0.00)
                    paragraph_format.right_indent = Inches(0.00)
                    
                    for run in paragraph.runs:
                        font = run.font
                        font.name="Arial"
                        font.size= Pt(8)
            for ii, cell in enumerate(table_cells[myNumColum*2:myNumColum*3]):
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # Adding styling        
                    set_cell_border(cell, bottom={"color": "#000000", "val": "single"})
            for cell in table_cells:
                table.cell=cell

            for cell in table.columns[0].cells:
                table.autofit = False
                table.allow_autofit = False
                cell.width=Inches(1.4)
                #table.autofit = True
                #table.allow_autofit = True
                for paragraph in cell.paragraphs:
                    ##pass
                    paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
            for i in range(myNumColum-1):
                for cell in table.columns[i+1].cells:
                    cell.width=Inches(float(scell_width))
                    #pass

            #table.columns[0].style=document.styles['itpvottabcol1']  
            
        print(f"time to generate itpvot is {time.time()-ts1}")
        
        add_page_number(document.sections[0].footer.paragraphs[0])
        return document


    if (bodysession=="A/10emsp"):
        #start = time.perf_counter()
        # Retrieving the number of records
        recordNumber = 0

        #bg - 
        myRecords=lstVotes
        myRecords1=lstVotes
        myRecords2=lstVotes
        for record in myRecords:
            recordNumber+=1
        

        # two columns display

        # Country list
        countries=[]     
        for datax in myRecords2[0]["votelist"]:
            countries.append(datax["memberstate"])

        # Adding the Header to the document

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
        
        # Marging of the document

        section.top_margin = Inches(0.81)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1) 

        # set some variables
        myRow = 0
        myCol = 0
        myNumColum=15
        myNumTable=0
        #myNumRow=206 # number of countries that have a vote
        myNumRow=len(countries)
        myIndex=0
        

        # generate the number of table
        myNumTable=math.ceil(recordNumber/myNumColum)

        # creation of the table
        table = document.add_table(rows=myNumRow+3, cols=myNumColum)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_cells=table._cells
        # write the memberstates in the first (0) column
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        myRow+=1
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        myRow+=1
        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
        
        for country in countries:
            myRow+=1
            table_cells[myRow*myNumColum].paragraphs[0].text=country
            
        myRow=0


        # Flip
        myFlip=0

        # set the header visible
        set_repeat_table_header(table.rows[0])
        set_repeat_table_header(table.rows[1])
        set_repeat_table_header(table.rows[2])

        
        myOneTimeExec=False
        vlCounter=0
        
        ts1=time.time()
        while myIndex<recordNumber:
            # for 1 row which is a header populate resolution numbers. Votes start from idx=3 and and with idx=195
            
            vlLen=len(myRecords1[myIndex]["votelist"])
            myRow=0
            
            
            if myRow == 0:
                if myOneTimeExec==False:
                    myResCounter=0
                    table_cells[myRow+myCol*myNumColum].paragraphs[0].text=myRecords1[myIndex]["docsymbol"][:12]+"-"
                    myOneTimeExec=True
                # adding two additional row
                
                if myFlip==0:
                    myFlip=1
                    #table_cells[myCol+1+myNflask runumColum].paragraphs[0].text=myRecords1[myIndex]["resnum"]
                    paragraph=table_cells[myCol+1+myNumColum].paragraphs[0]
                    add_hyperlink(paragraph,myRecords1[myIndex]["resnum"],"http://undocs.org/"+myRecords1[myIndex]["docsymbol"])
                    myResCounter+=1   
                else:
                    myFlip=0
                    #table_cells[myCol+1].paragraphs[0].text=myRecords1[myIndex]["resnum"]
                    paragraph=table_cells[myCol+1].paragraphs[0]
                    add_hyperlink(paragraph,myRecords1[myIndex]["resnum"],"http://undocs.org/"+myRecords1[myIndex]["docsymbol"])
                    myResCounter+=1
                    

                myRow+=1
                myRow+=1
                myRow+=1
            #write 1 column of votes
            
            for myVoteList in myRecords1[myIndex]["votelist"]: 
                table_cells[(myRow)*(myNumColum)+myCol+1].paragraphs[0].text=myVoteList["vote"]
                myRow+=1
                vlCounter+=1    
                if vlCounter==vlLen and myCol+1==14:
                    # check if we need to create a new table
                        myCol=0
                        myRow=0
                        myIndex+=1
                        vlCounter=0
                        #create a new table
                        p=document.add_paragraph()
                        run = p.add_run()
                        myBreak = run.add_break(WD_BREAK.PAGE)
                        table = document.add_table(rows=myNumRow+3, cols=myNumColum)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table_cells=table._cells
                        # write the memberstates in the first (0) column
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
                        myRow+=1
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""
                        myRow+=1
                        table_cells[myCol+myRow*myNumColum].paragraphs[0].text=""


                        # set the header visible
                        set_repeat_table_header(table.rows[0])
                        set_repeat_table_header(table.rows[1])
                        set_repeat_table_header(table.rows[2])

                        
                        for country in countries:
                            myRow+=1
                            table_cells[myRow*myNumColum].paragraphs[0].text=country
                        
                        myOneTimeExec=False

                    #write next column 
                elif vlCounter==vlLen:    
                        myCol+=1
                        myRow=0
                        myIndex+=1
                        vlCounter=0
                        
                    
            
        stl_itp_vot_cell= document.styles.add_style('itpvotcell', WD_STYLE_TYPE.PARAGRAPH)
        stl_itp_vot_cell_font=stl_itp_vot_cell.font
        stl_itp_vot_cell_font.name = 'Arial'
        stl_itp_vot_cell_font.size = Pt(8)
        stl_itp_vot_cell_font.bold = False
        #pf_stl_itp_vot_cell = stl_itp_vot_cell.paragraph_format

        for table in document.tables:
            table.autofit = False
            table.allow_autofit = False
            table_cells=table._cells
            for cell in table_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.style=stl_itp_vot_cell
                    ##pass
                    # Adding styling
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.line_spacing =  1
                    paragraph_format.left_indent = Inches(0.00)
                    paragraph_format.right_indent = Inches(0.00)
                    
                    for run in paragraph.runs:
                        font = run.font
                        font.name="Arial"
                        font.size= Pt(8)
            for ii, cell in enumerate(table_cells[30:45]):
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # Adding styling        
                    set_cell_border(cell, bottom={"color": "#000000", "val": "single"})
            for cell in table_cells:
                table.cell=cell

            for cell in table.columns[0].cells:
                cell.width=Inches(1.4)
                for paragraph in cell.paragraphs:
                    ##pass
                    paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
            for i in range(myNumColum-1):
                for cell in table.columns[i+1].cells:
                    cell.width=Inches(0.2)
            
            #table.columns[0].style=document.styles['itpvottabcol1']  
            
        print(f"time to generate itpvot is {time.time()-ts1}")
        
        add_page_number(document.sections[0].footer.paragraphs[0])
        return document


def generateWordDocITPREPS(paramTitle,paramSubTitle,bodysession,paramSection,paramNameFileOutput):
    
    # Function to keep the header visible for the table

    def set_repeat_table_header(row):

        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    # Setting some Variables

    myMongoURI=Config.connect_string
    #myClient = MongoClient(myMongoURI)
    #myDatabase=myClient.undlFiles
    client_dev_atlas=MongoClient(Config.connect_string_dev_atlas, tlsCAFile=certifi.where())
    myDatabase=client_dev_atlas.itpp
    myCollection=myDatabase['itp_sample_output_copy']
    myTitle=paramTitle
    mySubTitle1=paramSubTitle
    setOfData=myCollection.find({'bodysession': bodysession,'section': paramSection})
    
    # Creation of the word document
    document = Document()
    
    ################## HEADER ###############################################
    
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    
    # Font settings
    
    font = new_heading_style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Adding the header to the document
    
    header=document.sections[0].header
    
    ################## SUBHEADER ###############################################
    
    new_sub_heading_style = styles.add_style('New sub Heading', WD_STYLE_TYPE.PARAGRAPH)
   
    # Font settings
    pformat= new_sub_heading_style.paragraph_format
    pformat.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pformat.space_before = Pt(0)
    pformat.space_after = Pt(0)
    
    font = new_sub_heading_style.font
    font.name = 'Arial'
    font.size = Pt(7)
    font.italic = True
    font.color.rgb = RGBColor(0, 0, 0)

    ################## COMMITTEE ###############################################
    
    committee_style = styles.add_style('committee', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font settings
    pformat= committee_style.paragraph_format
    pformat.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pformat.space_after=Inches(0)

    font = committee_style.font
    font.name = 'Arial'
    font.size = Pt(9)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    ################## SUBJECT ###############################################
    
    subject_style = styles.add_style('subject', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font settings
    pformat= subject_style.paragraph_format
    pformat.alignment=WD_ALIGN_PARAGRAPH.LEFT
    pformat.space_after=Inches(0.05)
    pformat.first_line_indent=Inches(-0.07)

    
    font = subject_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.color.rgb = RGBColor(0, 0, 0)

    ################## DOCSYMBOL ###############################################
    
    docsymbol_style = styles.add_style('docsymbol', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font settings
    pformat= docsymbol_style.paragraph_format
    pformat.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    pformat.space_after=Inches(0.05)

    
    font = docsymbol_style.font
    font.name = 'Arial'
    font.size = Pt(8)
    font.color.rgb = RGBColor(0, 0, 0)

    ################## PSPACE ###############################################
    
    pspace_style = styles.add_style('pspace', WD_STYLE_TYPE.PARAGRAPH)
    
    # Font settings
    pformat= pspace_style.paragraph_format
    pformat.space_after=Inches(0)
    pformat.space_before=Inches(0)

    ################## WRITING THE DOCUMENT ###############################################
    
    # Header Generation
    
    p=header.add_paragraph(myTitle.upper(), style='New Heading')
    #p=header.add_paragraph(myTitle, style='New Heading')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    run=p.add_run("\n")

    # First line generation

    p=document.add_paragraph("The Committees discuss agenda items allocated to them by the Assembly.",style='New sub Heading')
    p=document.add_paragraph("The Committees report to the Assembly on the discussions held on each",style='New sub Heading')
    p=document.add_paragraph("item and forward recommendations for action in plenary.",style='New sub Heading') 

    # Definition of the column names
    myRecords=setOfData
    
    def getDocSymbols(myList):
        if len(myList)==1:
            add_hyperlink(row.cells[1].paragraphs[0],myList[0], Config.url_prefix+myList[0])

        else:
            myFinalResult=[]
            # storage for the root(s)
            myRoot=[]
            # storage for the Add(s)
            myAdd=[]
            # Storage for the Corr(s)
            myCorr=[]

            for myListElem in myList:
                recup=myListElem.split("/")

                # retrieving the root(s)
                if len(recup)==3 :
                    myRoot.append(myListElem)

                # retrieving the add(s)
                if len(recup)==4 and recup[3].startswith("Add"):
                    myAdd.append(myListElem)

                # retrieving the Corr(s)
                if len(recup)==4 and recup[3].startswith("Corr"):
                    myCorr.append(myListElem)

            myLink=""
            firstAdd=True
            firstCorr=True
            myLen=0 
            for root in myRoot:
                add_hyperlink(row.cells[1].paragraphs[0],root, Config.url_prefix+root)

                #adding of the Add if there is some values
                if len(myAdd)!=0:
                   
                   for add in myAdd:
                        if add.startswith(root):
                            recup1=add.split("/")
                            if firstAdd==True:
                                # adding the '+' sign
                                myParagraph = row.cells[1].paragraphs[0]
                                myParagraph.add_run("+")
                                if len(recup1[3])==5:
                                    add_hyperlink(row.cells[1].paragraphs[0],"Add."+add[-1], Config.url_prefix+add)
                                if len(recup1[3])==6: 
                                    xxx=add[-2]+add[-1]
                                    add_hyperlink(row.cells[1].paragraphs[0],"Add."+xxx, Config.url_prefix+add)
                                firstAdd=False
                            else:
                                myParagraph = row.cells[1].paragraphs[0]
                                myParagraph.add_run("-")
                                if len(recup1[3])==5:    
                                    add_hyperlink(row.cells[1].paragraphs[0],add[-1], Config.url_prefix+add)
                                if len(recup1[3])==6:    
                                    xxx=add[-2]+add[-1]
                                    add_hyperlink(row.cells[1].paragraphs[0],xxx, Config.url_prefix+add)    

                # adding of the Corr if there is some values
                if len(myCorr)!=0:
                   for corr in myCorr:
                        if corr.startswith(root):                
                            # adding the '+' sign
                            myParagraph = row.cells[1].paragraphs[0]
                            myParagraph.add_run(", ")
                            add_hyperlink(row.cells[1].paragraphs[0],corr, Config.url_prefix+corr)

                # adding sign , when you have a new record
                myLen+=1
                if myLen<len(myRoot):
                    myParagraph = row.cells[1].paragraphs[0]
                    myParagraph.add_run(", ") 

    for record in myRecords:
        
        # create a new paragraph
        p=document.add_paragraph("",style="pspace")

        # creation of the table
        table = document.add_table(rows=1, cols=2)

        # set the header visible
        set_repeat_table_header(table.rows[0])

        # paragraph=table.rows[0].cells[0].paragraphs[0]
        paragraph=table.rows[0].cells[0].merge(table.rows[0].cells[-1]).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text= record["committee"]
        paragraph.style = document.styles['committee']

        # writing the subjects
        for subject in record["subjects"]:

            # add a new row inside the table
            row = table.add_row()

            # add the content for the subject
            row.cells[0].paragraphs[0].text=subject['subject']
            row.cells[0].paragraphs[0].style = document.styles['subject']

            # add the content for the document symbol
            getDocSymbols(subject["docsymbols"])
            row.cells[1].paragraphs[0].style = document.styles['docsymbol']


        widths = (Inches(15), Inches(5.00))
        for row in table.rows:
            # sizing the rows
            row.height = Cm(0.2)
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(2) #3
            section.right_margin = Cm(2) #1.5

    # Save the document
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document